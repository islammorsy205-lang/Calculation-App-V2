# backprop_master.py

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from math_solver import get_prop_allowable, get_scaffold_allowable
except ImportError:
    st.error("⚠️ لم يتم العثور على math_solver.py. برجاء التأكد من مسار الملفات.")
    def get_prop_allowable(*args): return 20.0
    def get_scaffold_allowable(*args): return 30.0

def calculate_backprop_loads(fresh_slab, existing_slabs):
    results = []
    W_attacking = fresh_slab['total_load']
    results.append({'level': 'Fresh Slab', 'attacking': W_attacking, 'capacity': 0, 'transferred': W_attacking})
    
    current_P = W_attacking
    for i, slab in enumerate(existing_slabs):
        avail_cap = slab['capacity']
        
        absorbed = min(current_P, avail_cap)
        current_P = max(0, current_P - avail_cap)
        
        results.append({
            'level': f'Existing Slab {i+1}', 
            'attacking': results[-1]['transferred'], 
            'capacity': avail_cap, 
            'transferred': current_P,
            'slab_data': slab
        })
        
        if current_P <= 0:
            break
            
    return results, current_P

def plot_backprop_system(results):
    num_levels = len(results)
    fig, ax = plt.subplots(figsize=(10, num_levels * 1.5))
    
    y_pos = np.arange(num_levels, 0, -1) * 2
    
    for i, res in enumerate(results):
        y = y_pos[i]
        color = 'gray' if 'Existing' in res['level'] else 'blue'
        rect = plt.Rectangle((1, y-0.2), 8, 0.4, facecolor=color, alpha=0.5, edgecolor='black', linewidth=2)
        ax.add_patch(rect)
        ax.text(5, y, res['level'], ha='center', va='center', fontsize=12, fontweight='bold', color='white')
        
        if i < num_levels - 1 and res['transferred'] > 0:
            next_y = y_pos[i+1]
            ax.annotate('', xy=(5, next_y+0.2), xytext=(5, y-0.2),
                        arrowprops=dict(facecolor='red', shrink=0.05, width=4, headwidth=10))
            ax.text(5.2, (y + next_y)/2, f"{res['transferred']:.2f} kN/m²", color='red', fontsize=11, fontweight='bold')
            
            ax.plot([3, 3], [y-0.2, next_y+0.2], color='black', linewidth=3)
            ax.plot([7, 7], [y-0.2, next_y+0.2], color='black', linewidth=3)
            
    ax.set_xlim(0, 10)
    ax.set_ylim(0, max(y_pos) + 1)
    ax.axis('off')
    plt.title("Back-propping Load Transfer Diagram", fontsize=14, fontweight='bold')
    
    plt.tight_layout()
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_buf.seek(0)
    return img_buf

def generate_backprop_report(fresh_slab, results, grid_data, img_buf, ref_code):
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
        doc.add_page_break()
    else:
        doc = Document()
        
    def add_p(text, bold=False, underline=False, color=None, size=12, indent=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        if indent > 0:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.underline = underline
        if color:
            r.font.color.rgb = color
        return p

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CALCULATION SHEET FOR RE-PROPPING (BACK-PROPPING)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    
    p_code = doc.add_paragraph()
    p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_code = p_code.add_run("="*50 + f"\nCode Reference: {ref_code}")
    r_code.font.name = 'Arial'
    r_code.font.bold = True
    
    # ---------------------------------------------------------
    # Design Loads Section (Matching the Image)
    # ---------------------------------------------------------
    add_p("Design Loads for Slabs Back-Propping", bold=True, underline=True, color=RGBColor(192, 0, 0), size=14)
    
    add_p("A. Dead load:", indent=1)
    add_p(f"-  O.W of Concrete (Concrete density = {fresh_slab['gamma_c']} KN/m³)", indent=1)
    add_p(f"-  O.W of Formwork = {fresh_slab['FW']:.2f} KN/m²", indent=1)
    
    add_p("B. Live load:", indent=1)
    add_p(f"-  Live load = {fresh_slab['LL']:.2f} KN/m²", indent=1)
    
    add_p("\nLoad Acting on \"1st Level Slabs\" while Casting Level \"2nd Slabs\".", bold=True, underline=True)
    add_p(f"-  Slabs: {fresh_slab['ts_fresh'] * 1000:.0f} mm.", indent=1)
    
    W_slab_str = f"W Slab (KN/m²) = O.W of Slab + Live Load + O.W of Formwork"
    add_p(W_slab_str, bold=True, indent=1)
    calc_str = f"               = {fresh_slab['gamma_c']}X{fresh_slab['ts_fresh']:.2f} + {fresh_slab['LL']:.2f} + {fresh_slab['FW']:.2f} = {fresh_slab['total_load']:.2f} KN/m²"
    add_p(calc_str, bold=True, indent=1)
    
    if fresh_slab['has_drop']:
        add_p(f"\n-  Drops: {fresh_slab['ts_drop'] * 1000:.0f} mm.", indent=1)
        W_drop_str = f"W Drop (KN/m²) = O.W of Drop + Live Load + O.W of Formwork"
        add_p(W_drop_str, bold=True, indent=1)
        calc_drop = f"               = {fresh_slab['gamma_c']}X{fresh_slab['ts_drop']:.2f} + {fresh_slab['LL']:.2f} + {fresh_slab['FW']:.2f} = {fresh_slab['drop_load']:.2f} KN/m²"
        add_p(calc_drop, bold=True, indent=1)

    # ---------------------------------------------------------
    # Iterations Section
    # ---------------------------------------------------------
    for i, res in enumerate(results):
        if 'Existing' in res['level']:
            slab_data = res['slab_data']
            add_p(f"\n❖ Characteristic Surface Loads for Critical Zone ({res['level']}):", bold=True, underline=True)
            add_p(f"-  Super-imposed Dead Load (SDL) = {slab_data['sidl']:.2f} KN/m²", indent=1)
            add_p(f"-  Live Load (L.L) = {slab_data['ll']:.2f} KN/m²", indent=1)
            
            add_p(f"\nTotal un-Factored Load (W) = {slab_data['sidl']:.2f} + {slab_data['ll']:.2f} = {slab_data['unfactored']:.2f} KN/m²")
            add_p(f"Assume the concrete reach {slab_data['strength']:.0f}% from its strength.")
            add_p(f"Therefore: - Total resisting load = {slab_data['unfactored']:.2f} x {slab_data['strength']/100:.2f} = {res['capacity']:.2f} KN/m²")
            
            add_p(f"\nRe-Shoring Check for the System loaded on {res['level']}", bold=True)
            add_p(f"➢ Total Re-Shoring Loads from upper level = {res['attacking']:.2f} KN/m²", indent=1)
            
            add_p("Therefore; -", bold=True)
            add_p(f"The Transferred Loads to the Lower Level = {res['attacking']:.2f} - {res['capacity']:.2f} = {res['transferred']:.2f} KN/m²")

    # ---------------------------------------------------------
    # Shoring System Design
    # ---------------------------------------------------------
    if grid_data['needed']:
        add_p("\nShoring System Design Check:", bold=True)
        area_str = f"Max. Loaded Area \"Back Propped area\" = {grid_data['grid_x']:.2f} x {grid_data['grid_y']:.2f} = {grid_data['area_act']:.2f} m²"
        add_p(area_str, indent=1)
        
        load_leg = grid_data['area_act'] * grid_data['max_P']
        check_txt = f"Area Load on one leg of {grid_data['system']} = {grid_data['area_act']:.2f} x {grid_data['max_P']:.2f} = {load_leg:.3f} KN < {grid_data['allowable']:.2f} KN"
        
        p_check = doc.add_paragraph()
        p_check.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_check.paragraph_format.line_spacing = 1.5
        p_check.paragraph_format.left_indent = Cm(1)
        r_c1 = p_check.add_run(check_txt)
        r_c1.font.name = 'Arial'
        r_c1.font.size = Pt(12)
        
        r_res = p_check.add_run("   SAFE" if load_leg <= grid_data['allowable'] else "   UNSAFE")
        r_res.font.name = 'Arial'
        r_res.font.size = Pt(12)
        r_res.font.bold = True
        r_res.font.color.rgb = RGBColor(0, 128, 0) if load_leg <= grid_data['allowable'] else RGBColor(255, 0, 0)

    doc.add_page_break()
    add_p("Load Path Diagram:", bold=True)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(io.BytesIO(img_buf.read()), width=Cm(15.0))
    
    out = io.BytesIO()
    doc.save(out)
    return out

def render_backprop_module(ref_code):
    st.markdown("## 🏗️ Slab Re-propping (Back-propping) Designer")
    st.info("💡 **Independent Module:** Calculates load propagation and required shoring grid. Detailed outputs are strictly aligned to the left with 1.5 line spacing.")
    
    LL_const = 1.50 if "BS" in ref_code else 2.40
    FW_load = 0.50
    st.success(f"**Code Detected:** {ref_code}  →  Construction Live Load = {LL_const} kN/m², Formwork = {FW_load} kN/m²")
    
    st.markdown("### 1. Fresh Slab Data (Top Slab)")
    c1, c2, c3 = st.columns(3)
    gamma_c = c1.number_input("Concrete Density (kN/m³)", value=25.0, step=0.5, key='bp_gamma')
    ts_fresh = c2.number_input("Fresh Slab Thickness (m)", value=0.28, step=0.01, key='bp_ts')
    has_drop = c3.toggle("Include Drop Panels?")
    ts_drop = 0.60
    if has_drop:
        ts_drop = st.number_input("Drop Panel Thickness (m)", value=0.60, step=0.05)
    
    W_fresh = (gamma_c * ts_fresh) + LL_const + FW_load
    st.info(f"**Total Fresh Slab Load = {W_fresh:.2f} kN/m²**")
    W_drop = 0
    if has_drop:
        W_drop = (gamma_c * ts_drop) + LL_const + FW_load
        st.info(f"**Total Fresh Drop Load = {W_drop:.2f} kN/m²**")
        
    fresh_data = {
        'gamma_c': gamma_c, 'ts_fresh': ts_fresh, 'LL': LL_const, 'FW': FW_load, 
        'total_load': W_drop if has_drop else W_fresh, 'has_drop': has_drop, 'ts_drop': ts_drop, 'drop_load': W_drop
    }
    
    st.markdown("### 2. Existing Slabs Data")
    num_slabs = st.number_input("Number of Existing Slabs below", min_value=1, max_value=5, value=2)
    
    existing_slabs = []
    for i in range(int(num_slabs)):
        with st.expander(f"Existing Slab {i+1} (Directly Below)"):
            cb2, cb3, cb4 = st.columns(3)
            LL_des = cb2.number_input("Live Load (L.L) (kN/m²)", value=2.50, step=0.5, key=f'bp_ll_{i}')
            SIDL_des = cb3.number_input("Super-imposed Dead Load (SDL) (kN/m²)", value=0.50, step=0.5, key=f'bp_sidl_{i}')
            strength = cb4.number_input("Strength Achieved (%)", value=80.0, max_value=100.0, step=5.0, key=f'bp_str_{i}')
            
            Total_Des = SIDL_des + LL_des
            Cap_i = Total_Des * (strength / 100.0)
            
            existing_slabs.append({
                'll': LL_des,
                'sidl': SIDL_des,
                'unfactored': Total_Des,
                'strength': strength,
                'capacity': Cap_i
            })
            st.caption(f"Calculated Capacity = {Cap_i:.2f} kN/m²")
            
    st.markdown("### 3. Shoring System Details")
    cg1, cg2, cg3 = st.columns(3)
    grid_x = cg1.number_input("Grid Spacing X (m)", value=2.30, step=0.1)
    grid_y = cg2.number_input("Grid Spacing Y (m)", value=2.10, step=0.1)
    area_act = grid_x * grid_y
    cg3.info(f"**Loaded Area = {area_act:.2f} m²**")
    
    c_s1, c_s2 = st.columns(2)
    sys_opts = ["Acrow Prop", "Cup-lock", "Ring-lock", "Shorebrace Frame"]
    t_nm = c_s1.selectbox("Shoring Type", sys_opts, key='bp_sys')
    
    t_al = 20.0
    if t_nm == "Shorebrace Frame":
        t_al = 54.00
    elif t_nm == "Cup-lock":
        subtype = c_s2.selectbox("Steel Grade", ["S355 (st.52)", "S235"], key="bp_cup_sub")
        unb = c_s2.number_input("Unbraced Length (m)", value=1.5, step=0.5, key="bp_cup_unb")
        t_al = get_scaffold_allowable("Cup-lock", subtype, unb)
    elif t_nm == "Ring-lock":
        subtype = c_s2.selectbox("Diameter", ["Ringlock 1.5\"", "Ringlock 2.0\""], key="bp_ring_sub")
        unb = c_s2.number_input("Unbraced Length (m)", value=1.5, step=0.5, key="bp_ring_unb")
        t_al = get_scaffold_allowable("Ring-lock", subtype, unb)
    elif t_nm == "Acrow Prop":
        from config import PROP_DB
        req_ext = c_s2.number_input("Prop Extension (m)", value=3.0, step=0.1, key="bp_prop_ext")
        valid_props = [k for k, v in PROP_DB.items() if v['min'] <= req_ext <= v['max']] if PROP_DB else ["Prop No.2", "Prop No.3"]
        if valid_props:
            sel_prop = c_s2.selectbox("Select Valid Prop", valid_props, key="bp_prop_sel")
            t_al = get_prop_allowable(sel_prop, req_ext, True)
        else:
            t_al = 0.0
            
    st.info(f"**Allowable Load per Leg = {t_al:.2f} kN**")
    
    st.markdown("---")
    if st.button("🚀 Calculate Load Propagation & Generate Report", type="primary", use_container_width=True):
        results, unabsorbed = calculate_backprop_loads(fresh_data, existing_slabs)
        
        max_transferred = max([r['transferred'] for r in results[:-1]]) 
        
        grid_data = {'needed': False}
        if max_transferred > 0 and t_al > 0:
            grid_data = {
                'needed': True,
                'system': t_nm,
                'allowable': t_al,
                'max_P': max_transferred,
                'grid_x': grid_x,
                'grid_y': grid_y,
                'area_act': area_act
            }
            
            load_on_leg = area_act * max_transferred
            if load_on_leg <= t_al:
                st.success(f"✅ System is SAFE. Load on leg = {load_on_leg:.2f} kN < {t_al:.2f} kN")
            else:
                st.error(f"❌ System is UNSAFE! Load on leg = {load_on_leg:.2f} kN > {t_al:.2f} kN. Reduce grid spacing.")
        elif max_transferred <= 0:
            st.success("🎉 Existing slabs can safely carry the load! No Back-propping required.")
            
        img_buf = plot_backprop_system(results)
        st.image(img_buf, use_container_width=True)
        
        docx_out = generate_backprop_report(fresh_data, results, grid_data, img_buf, ref_code)
        
        st.download_button("⬇️ Download Detailed Back-propping Calculation Sheet", 
                           data=docx_out.getvalue(), 
                           file_name="Back_Propping_Detailed_Report.docx", 
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
