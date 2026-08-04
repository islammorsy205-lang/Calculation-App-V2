# inclined_master.py

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    from config import SECTIONS_DB, STRUTS_DB
except ImportError:
    st.error("⚠️ برجاء التأكد من وجود ملف config.py")
    SECTIONS_DB = {}
    STRUTS_DB = {}

def solve_inclined_fea(nodes, elements, gravity_loads):
    num_nodes = len(nodes)
    NDOF = num_nodes * 3
    K = np.zeros((NDOF, NDOF))
    F = np.zeros(NDOF)
    
    for idx, el in enumerate(elements):
        n1, n2 = el['n1'], el['n2']
        x1, y1 = nodes[n1][0], nodes[n1][1]
        x2, y2 = nodes[n2][0], nodes[n2][1]
        
        L = np.hypot(x2 - x1, y2 - y1)
        if L < 1e-5: continue
        
        c, s = (x2 - x1) / L, (y2 - y1) / L
        el['L'], el['c'], el['s'] = L, c, s
        
        E, A, I = el['E'], el['A'], el.get('I', 1e-6)
        
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        
        if el['type'] == 'truss':
            k_loc = np.zeros((6, 6))
            k_loc[0, 0] = k_loc[3, 3] = E * A / L
            k_loc[0, 3] = k_loc[3, 0] = -E * A / L
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0],
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L],
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2],
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
        k_glob = T.T @ k_loc @ T
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        
        for r in range(6):
            for col in range(6):
                K[dof_idx[r], dof_idx[col]] += k_glob[r, col]
                
    for gl in gravity_loads:
        el_idx = gl['mem_idx']
        W_gravity = gl['W'] 
        el = elements[el_idx]
        L, c, s = el['L'], el['c'], el['s']
        
        p_x = -W_gravity * s
        p_y = -W_gravity * c
        
        f_eq_loc = np.array([
            p_x * L / 2.0,
            p_y * L / 2.0,
            p_y * L**2 / 12.0,
            p_x * L / 2.0,
            p_y * L / 2.0,
            -p_y * L**2 / 12.0
        ])
        
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        f_eq_glob = T.T @ f_eq_loc
        
        n1, n2 = el['n1'], el['n2']
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        for r in range(6):
            F[dof_idx[r]] += f_eq_glob[r] 
            
    free_dof = []
    for i, n in enumerate(nodes):
        if not n[2]: free_dof.append(3*i)   
        if not n[3]: free_dof.append(3*i+1) 
        if not n[4]: free_dof.append(3*i+2) 
        
    K_ff = K[np.ix_(free_dof, free_dof)]
    F_f = F[free_dof]
    
    U = np.zeros(NDOF)
    try:
        U_f = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError:
        U_f = np.linalg.lstsq(K_ff, F_f, rcond=None)[0]
    
    U[free_dof] = U_f
    R_reactions = K @ U - F
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        c, s, L = el['c'], el['s'], el['L']
        E, A, I = el['E'], el['A'], el.get('I', 1e-6)
        
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        u_glob = U[dof_idx]
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        u_loc = T @ u_glob
        
        if el['type'] == 'truss':
            N = (E * A / L) * (u_loc[3] - u_loc[0])
            el['internal'] = {'N': N, 'V': np.zeros(11), 'M': np.zeros(11), 'x': np.linspace(0, L, 11)}
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0],
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L],
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2],
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            p_x, p_y = 0, 0
            for gl in gravity_loads:
                if elements[gl['mem_idx']] == el:
                    W_gravity = gl['W']
                    p_x = -W_gravity * s
                    p_y = -W_gravity * c
                    
            f_fixed_loc = np.array([
                -p_x * L / 2.0,
                -p_y * L / 2.0,
                -p_y * L**2 / 12.0,
                -p_x * L / 2.0,
                -p_y * L / 2.0,
                p_y * L**2 / 12.0
            ])
            
            f_end = k_loc @ u_loc + f_fixed_loc
            
            xs = np.linspace(0, L, 11)
            N_arr = -f_end[0] - p_x * xs
            V_arr = f_end[1] + p_y * xs
            M_arr = -f_end[2] + f_end[1] * xs + 0.5 * p_y * xs**2
            
            el['internal'] = {'N': N_arr, 'V': V_arr, 'M': M_arr, 'x': xs}
            
    return U, R_reactions, nodes, elements

def plot_inclined_system(nodes, elements, R_reactions):
    fig, axes = plt.subplots(2, 2, figsize=(16, 12))
    ax_geom, ax_n, ax_v, ax_m = axes.flatten()
    
    for ax in axes.flatten():
        ax.set_aspect('equal', adjustable='datalim')
        ax.grid(True, linestyle=':', alpha=0.5)
        
    def draw_supports(ax):
        for i, n in enumerate(nodes):
            x, y = n[0], n[1]
            if n[2] and n[3]: 
                ax.plot(x, y, marker='^', color='orange', markersize=12, zorder=5)
            elif not n[2] and n[3]: 
                ax.plot(x, y, marker='o', color='green', markersize=10, zorder=5)

    ax_geom.set_title("System Geometry & Supports", fontsize=12, fontweight='bold')
    for el in elements:
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        color = 'blue' if el['type'] == 'frame' else 'red'
        style = '-' if el['type'] == 'frame' else '--'
        ax_geom.plot([n1[0], n2[0]], [n1[1], n2[1]], color=color, linestyle=style, linewidth=2)
    draw_supports(ax_geom)
    
    ax_n.set_title("Axial Force Diagram (kN)", fontsize=12, fontweight='bold')
    for el in elements:
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        x1, y1, x2, y2 = n1[0], n1[1], n2[0], n2[1]
        c, s = el['c'], el['s']
        if el['type'] == 'truss':
            val = el['internal']['N']
            ax_n.plot([x1, x2], [y1, y2], color='red', linestyle='--', linewidth=1.5)
            ax_n.text((x1+x2)/2, (y1+y2)/2, f"{val:.1f}", color='red', fontsize=9, fontweight='bold')
        else:
            N_vals = el['internal']['N']
            xs = el['internal']['x']
            scale = 0.02
            for j in range(len(xs)-1):
                x_start = x1 + c * xs[j] - s * (N_vals[j]*scale)
                y_start = y1 + s * xs[j] + c * (N_vals[j]*scale)
                x_end = x1 + c * xs[j+1] - s * (N_vals[j+1]*scale)
                y_end = y1 + s * xs[j+1] + c * (N_vals[j+1]*scale)
                color = 'blue' if N_vals[j] < 0 else 'red'
                ax_n.plot([x_start, x_end], [y_start, y_end], color=color, linewidth=1.5)
            ax_n.plot([x1, x2], [y1, y2], color='black', linewidth=1)
            ax_n.text(x1 + c*xs[5], y1 + s*xs[5] + 0.5, f"Max: {np.max(np.abs(N_vals)):.1f}", color='black', fontsize=9)
    draw_supports(ax_n)

    ax_v.set_title("Shear Force Diagram (kN)", fontsize=12, fontweight='bold')
    for el in elements:
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        x1, y1, x2, y2 = n1[0], n1[1], n2[0], n2[1]
        c, s = el['c'], el['s']
        if el['type'] == 'frame':
            V_vals = el['internal']['V']
            xs = el['internal']['x']
            scale = 0.02
            for j in range(len(xs)-1):
                x_start = x1 + c * xs[j] - s * (V_vals[j]*scale)
                y_start = y1 + s * xs[j] + c * (V_vals[j]*scale)
                x_end = x1 + c * xs[j+1] - s * (V_vals[j+1]*scale)
                y_end = y1 + s * xs[j+1] + c * (V_vals[j+1]*scale)
                ax_v.plot([x_start, x_end], [y_start, y_end], color='purple', linewidth=1.5)
            ax_v.plot([x1, x2], [y1, y2], color='black', linewidth=1)
            ax_v.text(x1 + c*xs[5], y1 + s*xs[5] + 0.5, f"Max: {np.max(np.abs(V_vals)):.1f}", color='black', fontsize=9)
    draw_supports(ax_v)
    
    ax_m.set_title("Bending Moment Diagram (kN.m)", fontsize=12, fontweight='bold')
    for el in elements:
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        x1, y1, x2, y2 = n1[0], n1[1], n2[0], n2[1]
        c, s = el['c'], el['s']
        if el['type'] == 'frame':
            M_vals = el['internal']['M']
            xs = el['internal']['x']
            scale = 0.05
            for j in range(len(xs)-1):
                x_start = x1 + c * xs[j] - s * (M_vals[j]*scale)
                y_start = y1 + s * xs[j] + c * (M_vals[j]*scale)
                x_end = x1 + c * xs[j+1] - s * (M_vals[j+1]*scale)
                y_end = y1 + s * xs[j+1] + c * (M_vals[j+1]*scale)
                ax_m.plot([x_start, x_end], [y_start, y_end], color='green', linewidth=1.5)
            ax_m.plot([x1, x2], [y1, y2], color='black', linewidth=1)
            ax_m.text(x1 + c*xs[5], y1 + s*xs[5] + 0.5, f"Max: {np.max(np.abs(M_vals)):.1f}", color='black', fontsize=9)
    draw_supports(ax_m)
    
    plt.tight_layout()
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_buf.seek(0)
    return img_buf

def generate_inclined_report(sys_data):
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
        doc.add_page_break()
    else:
        doc = Document()
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CALCULATION SHEET FOR INCLINED FORMWORK SYSTEM")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    
    doc.add_paragraph("="*50)
    
    def add_line(text, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = bold
        
    def add_check(component, param, act, allw, unit):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        p.add_run(f"• Check {component} ({param}):\n").bold = True
        r_act = p.add_run(f"  Actual = {act:.2f} {unit}  <  Allowable = {allw:.2f} {unit}  ")
        res = "SAFE" if act <= allw else "UNSAFE"
        r_res = p.add_run(res)
        r_res.font.bold = True
        r_res.font.color.rgb = RGBColor(0, 128, 0) if res == "SAFE" else RGBColor(255, 0, 0)
        
    add_line(f"1. Geometry & Inputs:", bold=True)
    add_line(f"- Inclined Soldier Length = {sys_data['L_tot']:.2f} m")
    add_line(f"- Inclination Angle = {sys_data['angle']:.1f} degrees")
    add_line(f"- Total Gravity Load (W) = {sys_data['W']:.2f} kN/m")
    
    doc.add_paragraph()
    add_line(f"2. Safety Checks:", bold=True)
    
    inc_sec = sys_data['inc_sec']
    add_line(f"A. Inclined Soldier ({inc_sec})", bold=True)
    add_check("Moment", "M_max", sys_data['max_M_inc'], SECTIONS_DB.get(inc_sec, {}).get('Mall', 999), "kN.m")
    add_check("Shear", "V_max", sys_data['max_V_inc'], SECTIONS_DB.get(inc_sec, {}).get('Qall', 999), "kN")
    
    base_sec = sys_data['base_sec']
    add_line(f"B. Horizontal Base Soldier ({base_sec})", bold=True)
    add_check("Moment", "M_max", sys_data['max_M_base'], SECTIONS_DB.get(base_sec, {}).get('Mall', 999), "kN.m")
    add_check("Shear", "V_max", sys_data['max_V_base'], SECTIONS_DB.get(base_sec, {}).get('Qall', 999), "kN")
    
    add_line("C. Push-Pull Struts (Axial Force)", bold=True)
    for idx, st_val in enumerate(sys_data['struts_res']):
        allow = STRUTS_DB.get(st_val['type'], {}).get('allow', 999) 
        add_check(f"Strut {idx+1} ({st_val['type']})", "N_max", st_val['N'], allow, "kN")
        
    doc.add_page_break()
    add_line("3. Analysis Diagrams", bold=True)
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(io.BytesIO(sys_data['img_bytes'].read()), width=Cm(16.5))
    
    out = io.BytesIO()
    doc.save(out)
    return out

def render_inclined_module():
    st.markdown("## 📐 Inclined Formwork System (Advanced 2D Frame)")
    st.info("💡 **Independent Module:** This system solves inclined setups with Roller & Hinged supports, computing Gravity Loads precisely without affecting standard modules.")
    
    c1, c2, c3 = st.columns(3)
    with c1:
        angle_deg = st.number_input("Inclination Angle (Degrees, < 90)", value=60.0, step=5.0)
        angle_rad = np.radians(angle_deg)
    with c2:
        W_load = st.number_input("Vertical Gravity Load (kN/m)", value=15.0, step=1.0)
    with c3:
        num_struts = st.number_input("Number of Push-Pulls", min_value=1, max_value=5, value=2, step=1)
        
    st.markdown("---")
    st.markdown("#### 🪵 Sections & Segments Setup")
    
    col_inc, col_base = st.columns(2)
    
    with col_inc:
        st.markdown("**1. Inclined Soldier**")
        inc_sec = st.selectbox("Profile (Inclined)", list(SECTIONS_DB.keys()) if SECTIONS_DB else ["Soldier U100"])
        L_segs = []
        for j in range(int(num_struts)):
            L_segs.append(st.number_input(f"Segment L{j+1} (Distance to Strut {j+1}) (m)", value=2.0, step=0.5, key=f"L_seg_{j}"))
        L_rem = st.number_input("Remaining Length (Cantilever Top) (m)", value=1.0, step=0.5)
        L_tot = sum(L_segs) + L_rem
        st.success(f"Total Inclined Length = {L_tot:.2f} m")
        
    with col_base:
        st.markdown("**2. Horizontal Base Soldier**")
        base_sec = st.selectbox("Profile (Base)", list(SECTIONS_DB.keys()) if SECTIONS_DB else ["Soldier U100"])
        X_segs = []
        strut_types = []
        for j in range(int(num_struts)):
            X_segs.append(st.number_input(f"Segment X{j+1} (Distance to Base {j+1}) (m)", value=1.5, step=0.5, key=f"X_seg_{j}"))
            strut_types.append(st.selectbox(f"Strut {j+1} Profile", list(STRUTS_DB.keys()) if STRUTS_DB else ["PPH601"], key=f"st_type_{j}"))
        X_rem = st.number_input("Remaining Base Length (m)", value=0.5, step=0.5)
        X_tot = sum(X_segs) + X_rem
        st.success(f"Total Base Length = {X_tot:.2f} m")
        
    st.markdown("---")
    if st.button("🚀 Run Advanced FEA & Generate Report", type="primary", use_container_width=True):
        with st.spinner("Building Stiffness Matrix & Solving..."):
            
            nodes = []
            nodes.append([0.0, 0.0, False, True, False]) 
            
            L_cum = 0.0
            inc_node_indices = [0]
            for L_seg in L_segs:
                L_cum += L_seg
                x, y = L_cum * np.cos(angle_rad), L_cum * np.sin(angle_rad)
                nodes.append([x, y, False, False, False])
                inc_node_indices.append(len(nodes)-1)
            
            if L_rem > 0:
                L_cum += L_rem
                x, y = L_cum * np.cos(angle_rad), L_cum * np.sin(angle_rad)
                nodes.append([x, y, False, False, False])
                inc_node_indices.append(len(nodes)-1)
                
            X_cum = 0.0
            base_node_indices = [0]
            for X_seg in X_segs:
                X_cum += X_seg
                nodes.append([X_cum, 0.0, True, True, False])
                base_node_indices.append(len(nodes)-1)
                
            if X_rem > 0:
                X_cum += X_rem
                nodes.append([X_cum, 0.0, False, False, False]) 
                base_node_indices.append(len(nodes)-1)
                
            elements = []
            gravity_loads = []
            E_st = 210000000.0 
            
            inc_props = SECTIONS_DB.get(inc_sec, {'E': E_st, 'A': 0.00343, 'I': 0.00000122})
            for j in range(len(inc_node_indices)-1):
                elements.append({
                    'type': 'frame', 'sec': inc_sec,
                    'n1': inc_node_indices[j], 'n2': inc_node_indices[j+1],
                    'E': inc_props.get('E', E_st), 'A': inc_props.get('A', 0.00343), 'I': inc_props.get('I', 0.00000122)
                })
                gravity_loads.append({'mem_idx': len(elements)-1, 'W': W_load})
                
            base_props = SECTIONS_DB.get(base_sec, {'E': E_st, 'A': 0.00343, 'I': 0.00000122})
            for j in range(len(base_node_indices)-1):
                elements.append({
                    'type': 'frame', 'sec': base_sec,
                    'n1': base_node_indices[j], 'n2': base_node_indices[j+1],
                    'E': base_props.get('E', E_st), 'A': base_props.get('A', 0.00343), 'I': base_props.get('I', 0.00000122)
                })
                
            struts_results_placeholder = []
            for j in range(int(num_struts)):
                n_inc = inc_node_indices[j+1]
                n_base = base_node_indices[j+1]
                st_type = strut_types[j]
                st_props = STRUTS_DB.get(st_type, {'A': 0.001}) 
                elements.append({
                    'type': 'truss', 'sec': st_type,
                    'n1': n_base, 'n2': n_inc,
                    'E': E_st, 'A': st_props.get('A', 0.001)
                })
                struts_results_placeholder.append({'idx': len(elements)-1, 'type': st_type})
                
            U, R, final_nodes, final_elements = solve_inclined_fea(nodes, elements, gravity_loads)
            
            img_buf = plot_inclined_system(final_nodes, final_elements, R)
            st.image(img_buf, use_container_width=True)
            
            max_M_inc, max_V_inc = 0, 0
            max_M_base, max_V_base = 0, 0
            
            for el in final_elements:
                if el['type'] == 'frame':
                    if el['sec'] == inc_sec:
                        max_M_inc = max(max_M_inc, np.max(np.abs(el['internal']['M'])))
                        max_V_inc = max(max_V_inc, np.max(np.abs(el['internal']['V'])))
                    elif el['sec'] == base_sec:
                        max_M_base = max(max_M_base, np.max(np.abs(el['internal']['M'])))
                        max_V_base = max(max_V_base, np.max(np.abs(el['internal']['V'])))
            
            for st_res in struts_results_placeholder:
                el = final_elements[st_res['idx']]
                st_res['N'] = abs(el['internal']['N'])
                
            sys_data = {
                'L_tot': L_tot,
                'angle': angle_deg,
                'W': W_load,
                'inc_sec': inc_sec,
                'base_sec': base_sec,
                'max_M_inc': max_M_inc,
                'max_V_inc': max_V_inc,
                'max_M_base': max_M_base,
                'max_V_base': max_V_base,
                'struts_res': struts_results_placeholder,
                'img_bytes': img_buf
            }
            
            docx_out = generate_inclined_report(sys_data)
            
            st.success("✅ Analysis Complete!")
            st.download_button("⬇️ Download Inclined System Calculation Sheet", 
                               data=docx_out.getvalue(), 
                               file_name="Inclined_System_Report.docx", 
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
