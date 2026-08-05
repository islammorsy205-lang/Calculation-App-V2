# bridge_master.py

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================================================
# 0. Database Loading
# =========================================================
try:
    from config import SECTIONS_DB
except ImportError:
    st.error("⚠️ برجاء التأكد من وجود ملف config.py")
    SECTIONS_DB = {}

# =========================================================
# 1. Geometry & Mesh Generator (Parametric Truss Bridge)
# =========================================================
def build_bridge_mesh(num_bays, bay_length, depth, top_sec, bot_sec, diag_sec, vert_sec, w_load):
    nodes = []
    elements = []
    nodal_loads = []
    supports_list = []
    
    # 1. Generate Nodes
    # Bottom Chord Nodes (y=0)
    for i in range(num_bays + 1):
        nodes.append([i * bay_length, 0.0])
    
    # Top Chord Nodes (y=depth)
    for i in range(num_bays):
        nodes.append([(i + 0.5) * bay_length, depth])
        
    bottom_nodes_count = num_bays + 1
    
    # Supports (Simple span: Hinge at start, Roller at end)
    supports_list.append({'node': 0, 'type': 'Hinged', 'angle': 0.0})
    supports_list.append({'node': num_bays, 'type': 'Roller', 'angle': 0.0})
    
    # 2. Generate Elements
    def get_props(sec_name):
        return SECTIONS_DB.get(sec_name, {'E': 2100.0, 'A': 20.0, 'I': 412.0})

    # Bottom Chords
    props = get_props(bot_sec)
    for i in range(num_bays):
        elements.append({
            'type': 'frame', 'group': 'Bottom Chord', 'sec': bot_sec,
            'n1': i, 'n2': i + 1,
            'px1': 0, 'py1': -w_load, 'px2': 0, 'py2': -w_load, # Load applied to bottom chord
            'E': props.get('E', 2100.0) * 10000.0, 
            'A': props.get('A', 20.0) / 10000.0, 
            'I': props.get('I', 412.0) / 100000000.0
        })
        
    # Top Chords
    props = get_props(top_sec)
    for i in range(num_bays - 1):
        elements.append({
            'type': 'frame', 'group': 'Top Chord', 'sec': top_sec,
            'n1': bottom_nodes_count + i, 'n2': bottom_nodes_count + i + 1,
            'px1': 0, 'py1': 0, 'px2': 0, 'py2': 0,
            'E': props.get('E', 2100.0) * 10000.0, 
            'A': props.get('A', 20.0) / 10000.0, 
            'I': props.get('I', 412.0) / 100000000.0
        })

    # Diagonals & Verticals
    props_diag = get_props(diag_sec)
    props_vert = get_props(vert_sec)
    
    for i in range(num_bays):
        # Diagonal /
        elements.append({
            'type': 'frame', 'group': 'Diagonal', 'sec': diag_sec,
            'n1': i, 'n2': bottom_nodes_count + i,
            'px1': 0, 'py1': 0, 'px2': 0, 'py2': 0,
            'E': props_diag.get('E', 2100.0) * 10000.0, 
            'A': props_diag.get('A', 20.0) / 10000.0, 
            'I': props_diag.get('I', 412.0) / 100000000.0
        })
        # Diagonal \
        elements.append({
            'type': 'frame', 'group': 'Diagonal', 'sec': diag_sec,
            'n1': bottom_nodes_count + i, 'n2': i + 1,
            'px1': 0, 'py1': 0, 'px2': 0, 'py2': 0,
            'E': props_diag.get('E', 2100.0) * 10000.0, 
            'A': props_diag.get('A', 20.0) / 10000.0, 
            'I': props_diag.get('I', 412.0) / 100000000.0
        })
        # Verticals
        if i > 0:
            elements.append({
                'type': 'frame', 'group': 'Vertical', 'sec': vert_sec,
                'n1': i, 'n2': bottom_nodes_count + i - 1, # Connecting to previous top node
                'px1': 0, 'py1': 0, 'px2': 0, 'py2': 0,
                'E': props_vert.get('E', 2100.0) * 10000.0, 
                'A': props_vert.get('A', 20.0) / 10000.0, 
                'I': props_vert.get('I', 412.0) / 100000000.0
            })
            elements.append({
                'type': 'frame', 'group': 'Vertical', 'sec': vert_sec,
                'n1': i, 'n2': bottom_nodes_count + i, # Connecting to current top node
                'px1': 0, 'py1': 0, 'px2': 0, 'py2': 0,
                'E': props_vert.get('E', 2100.0) * 10000.0, 
                'A': props_vert.get('A', 20.0) / 10000.0, 
                'I': props_vert.get('I', 412.0) / 100000000.0
            })

    L_tot = num_bays * bay_length
    display_nodes = set(range(len(nodes)))
    
    return nodes, elements, nodal_loads, L_tot, depth, display_nodes, supports_list

# =========================================================
# 2. FEA Solver (Exact match with existing engine)
# =========================================================
def solve_fea_engine(nodes, elements, nodal_loads, supports_list):
    num_nodes = len(nodes)
    NDOF = num_nodes * 3
    K = np.zeros((NDOF, NDOF))
    F = np.zeros(NDOF)
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        x1, y1 = nodes[n1][0], nodes[n1][1]
        x2, y2 = nodes[n2][0], nodes[n2][1]
        L = np.hypot(x2 - x1, y2 - y1)
        if L < 1e-5: continue
        c, s = (x2 - x1) / L, (y2 - y1) / L
        el['L'], el['c'], el['s'] = L, c, s
        
        E, A, I = el['E'], el['A'], el.get('I', 0.00005)
        
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        
        if el['type'] == 'truss':
            k_loc = np.zeros((6, 6))
            k_loc[0, 0] = E * A / L
            k_loc[3, 3] = E * A / L
            k_loc[0, 3] = -E * A / L
            k_loc[3, 0] = -E * A / L
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0],
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L],
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2],
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            px1, py1, px2, py2 = el['px1'], el['py1'], el['px2'], el['py2']
            f_eq_loc = np.array([
                (2*px1 + px2)*L/6.0,
                (7*py1 + 3*py2)*L/20.0,
                (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0,
                (3*py1 + 7*py2)*L/20.0,
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_eq_glob = T.T @ f_eq_loc
            dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
            for r in range(6): F[dof_idx[r]] += f_eq_glob[r]
            
        k_glob = T.T @ k_loc @ T
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        for r in range(6):
            for col in range(6):
                K[dof_idx[r], dof_idx[col]] += k_glob[r, col]
                
    for nl in nodal_loads:
        F[3*nl['node'] + 0] += nl['Fx']
        F[3*nl['node'] + 1] += nl['Fy']
            
    K_orig = K.copy()
    fixed_dofs = []
    K_penalty = 1e12
    
    for sup in supports_list:
        n = sup['node']
        t = sup['type']
        a = sup['angle']
        if t == 'Fixed':
            fixed_dofs.extend([3*n, 3*n+1, 3*n+2])
        elif t == 'Hinged':
            fixed_dofs.extend([3*n, 3*n+1])
        elif t == 'Roller':
            if abs(a % 180) < 1e-5: 
                fixed_dofs.append(3*n+1)
            elif abs((a - 90) % 180) < 1e-5: 
                fixed_dofs.append(3*n)
            else:
                rad = np.radians(a)
                nx, ny = -np.sin(rad), np.cos(rad) 
                K[3*n, 3*n] += K_penalty * nx**2
                K[3*n+1, 3*n+1] += K_penalty * ny**2
                K[3*n, 3*n+1] += K_penalty * nx * ny
                K[3*n+1, 3*n] += K_penalty * nx * ny

    free_dof = [i for i in range(NDOF) if i not in fixed_dofs]
        
    K_ff = K[np.ix_(free_dof, free_dof)]
    F_f = F[free_dof]
    
    U = np.zeros(NDOF)
    try:
        U_f = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError:
        U_f = np.linalg.lstsq(K_ff, F_f, rcond=None)[0]
    
    U[free_dof] = U_f
    R_reactions = K_orig @ U - F 
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        c, s, L = el['c'], el['s'], el['L']
        E, A, I = el['E'], el['A'], el.get('I', 0.00005)
        
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        u_glob = U[dof_idx]
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        u_loc = T @ u_glob
        el['internal'] = {'u_loc': u_loc}
        
        if el['type'] == 'truss':
            N_val = (E * A / L) * (u_loc[3] - u_loc[0])
            el['internal'].update({'N': [N_val, N_val], 'V': [0,0], 'M': [0,0], 'x': [0, L], 'v_rel': [0,0]})
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0],
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L],
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2],
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            px1, py1, px2, py2 = el['px1'], el['py1'], el['px2'], el['py2']
            f_eq_loc = np.array([
                (2*px1 + px2)*L/6.0,
                (7*py1 + 3*py2)*L/20.0,
                (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0,
                (3*py1 + 7*py2)*L/20.0,
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_end = k_loc @ u_loc - f_eq_loc
            
            xs = np.linspace(0, L, 51) 
            N_arr = np.zeros_like(xs)
            V_arr = np.zeros_like(xs)
            M_arr = np.zeros_like(xs)
            v_rel_arr = np.zeros_like(xs)
            
            v1, theta1 = u_loc[1], u_loc[2]
            v2, theta2 = u_loc[4], u_loc[5]
            w_avg = (py1 + py2) / 2.0 
            
            for i, x in enumerate(xs):
                N_arr[i] = -f_end[0] - (px1*x + (px2-px1)*x**2/(2*L))
                V_arr[i] = f_end[1] + (py1*x + (py2-py1)*x**2/(2*L))
                M_arr[i] = -f_end[2] + f_end[1]*x + py1*x**2/2.0 + (py2-py1)*x**3/(6*L)
                
                xi = x / L
                N1 = 1 - 3*xi**2 + 2*xi**3
                N2 = x * (1 - xi)**2
                N3 = 3*xi**2 - 2*xi**3
                N4 = x * (xi**2 - xi)
                
                v_shape = v1*N1 + theta1*N2 + v2*N3 + theta2*N4
                v_load = (w_avg * x**2 * (L - x)**2) / (24 * E * I) 
                v_tot = v_shape + v_load
                v_chord = v1 + xi * (v2 - v1) 
                v_rel_arr[i] = v_tot - v_chord 
                
            el['internal'].update({'N': N_arr, 'V': V_arr, 'M': M_arr, 'x': xs, 'v_rel': v_rel_arr})
            
    return U, R_reactions

# =========================================================
# 3. Plotting Engine (SAP2000 Style)
# =========================================================
def get_img_buf(fig):
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight', transparent=True)
    plt.close(fig)
    buf.seek(0)
    return buf

def draw_base_geometry(ax, nodes, elements, supports_list):
    for el in elements:
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='black', linestyle='-', linewidth=0.8, zorder=1)
        
    for sup in supports_list:
        n = sup['node']
        x, y = nodes[n][0], nodes[n][1]
        t = sup['type']
        
        # SAP2000 Limegreen supports
        if t == 'Hinged':
            h, w = 0.5, 0.4
            p1 = (x, y)
            p2 = (x + w/2, y - h)
            p3 = (x - w/2, y - h)
            poly = Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.5, zorder=5)
            ax.add_patch(poly)
            ax.plot([x - w, x + w], [y - h, y - h], color='limegreen', lw=2.0, zorder=4)
        elif t == 'Roller':
            h, w, r = 0.4, 0.3, 0.15
            p1 = (x, y)
            p2 = (x + w/2, y - h)
            p3 = (x - w/2, y - h)
            poly = Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.5, zorder=5)
            ax.add_patch(poly)
            circle = plt.Circle((x, y - h - r), r, facecolor='none', edgecolor='limegreen', lw=1.5, zorder=5)
            ax.add_patch(circle)
            ax.plot([x - w, x + w], [y - h - 2*r, y - h - 2*r], color='limegreen', lw=2.0, zorder=4)

def plot_bridge_diagrams(nodes, elements, R_reactions, scales, display_nodes, L_tot, supports_list):
    mpl.rcParams['font.family'] = 'sans-serif'
    mpl.rcParams['font.sans-serif'] = ['Arial']
    mpl.rcParams['font.size'] = 7
    
    figs_dict = {}
    
    # --- Force Diagrams Helper ---
    def create_force_diagram(val_key, scale, is_axial=False):
        fig, ax = plt.subplots(figsize=(10, 4))
        ax.set_aspect('equal', adjustable='datalim')
        ax.axis('off')
        
        draw_base_geometry(ax, nodes, elements, supports_list)
        
        for el in elements:
            n1, n2 = nodes[el['n1']], nodes[el['n2']]
            x1, y1 = n1[0], n1[1]
            x2, y2 = n2[0], n2[1]
            dx, dy = x2 - x1, y2 - y1
            L_s = np.hypot(dx, dy)
            if L_s < 1e-5: continue
            
            c, s = dx/L_s, dy/L_s
            
            if el['type'] == 'frame':
                xs_arr = el['internal']['x']
                vals_orig = el['internal'][val_key]
                
                # Plotting logic: +ve = Blue, -ve = Red
                plot_vals = -vals_orig if val_key != 'N' else vals_orig
                
                px_arr = x1 + c * xs_arr - s * plot_vals * scale
                py_arr = y1 + s * xs_arr + c * plot_vals * scale
                
                for k in range(len(px_arr)-1):
                    avg_v = (vals_orig[k] + vals_orig[k+1]) / 2.0
                    seg_color = 'blue' if avg_v >= 0 else 'red'
                    
                    # Fill
                    p1 = (x1 + c * xs_arr[k], y1 + s * xs_arr[k])
                    p2 = (px_arr[k], py_arr[k])
                    p3 = (px_arr[k+1], py_arr[k+1])
                    p4 = (x1 + c * xs_arr[k+1], y1 + s * xs_arr[k+1])
                    
                    ax.add_patch(Polygon([p1, p2, p3, p4], facecolor=seg_color, alpha=0.3, edgecolor='none', zorder=2))
                    ax.plot([px_arr[k], px_arr[k+1]], [py_arr[k], py_arr[k+1]], color=seg_color, linewidth=1.0, zorder=3)
                
                # Labels for Max Values
                max_v_idx = np.argmax(np.abs(vals_orig))
                max_v = vals_orig[max_v_idx]
                if abs(max_v) > 0.1:
                    ax.text(px_arr[max_v_idx], py_arr[max_v_idx], f"{abs(max_v):.1f}", color='black', fontsize=7, ha='center', va='center')

        return get_img_buf(fig)

    figs_dict['N'] = create_force_diagram('N', scales['N'], is_axial=True)
    figs_dict['V'] = create_force_diagram('V', scales['V'])
    figs_dict['M'] = create_force_diagram('M', scales['M'])

    return figs_dict

# =========================================================
# 4. Report Generator
# =========================================================
def generate_bridge_report(sys_data):
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
        doc.add_page_break()
    else:
        doc = Document()
        
    def force_ltr_left(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)
        
    def add_line(text, bold=False, color=None):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.bold = bold
        r.font.rtl = False
        if color: r.font.color.rgb = color
        
    p_title = doc.add_paragraph()
    force_ltr_left(p_title)
    run_title = p_title.add_run("CALCULATION SHEET FOR BRIDGE STRUCTURE (TRUSS)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.rtl = False
    
    add_line("="*50, bold=True)
    add_line(f"1. Geometry & Inputs:", bold=True)
    add_line(f"- Total Bridge Span = {sys_data['L_tot']:.2f} m")
    add_line(f"- Truss Depth = {sys_data['depth']:.2f} m")
    add_line(f"- Top Chord: {sys_data['top_sec']} | Bottom Chord: {sys_data['bot_sec']}")
    add_line(f"- Diagonals: {sys_data['diag_sec']} | Verticals: {sys_data['vert_sec']}")
    add_line(f"- Applied Uniform Load on Bottom Chord = {sys_data['w_load']} kN/m")
    
    doc.add_paragraph()
    add_line("2. Analysis Diagrams (SAP2000 Style):", bold=True)
    
    titles = {
        'N': "Axial Force Diagram (kN) [Blue = Tension, Red = Compression]",
        'V': "Shear Force Diagram (kN)",
        'M': "Bending Moment Diagram (kN.m)"
    }
    
    for key in ['N', 'V', 'M']:
        buf = sys_data['img_bufs'][key]
        buf.seek(0)
        
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(buf.read()), width=Cm(16.0))
        
        p_txt = doc.add_paragraph()
        p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_txt = p_txt.add_run(titles[key])
        r_txt.font.name = 'Arial'
        r_txt.font.size = Pt(10)
        r_txt.underline = True
        r_txt.font.rtl = False
        
        doc.add_page_break()
    
    out = io.BytesIO()
    doc.save(out)
    return out

# =========================================================
# 5. Main UI Module
# =========================================================
def render_bridge_module():
    st.markdown("## 🌉 Bridge Formwork & Structures (Advanced FEA)")
    st.info("💡 **Bridge Module:** Parametric generator for Modular Truss Bridges (Warren type).")
    
    if 'bridge_solved' not in st.session_state:
        st.session_state.bridge_solved = False
        
    st.markdown("#### ⚙️ 1. Bridge Geometry & Sections")
    c1, c2, c3 = st.columns(3)
    num_bays = c1.number_input("Number of Bays", min_value=2, max_value=20, value=6, step=1, on_change=lambda: st.session_state.update(bridge_solved=False))
    bay_length = c2.number_input("Bay Length (m)", value=3.0, step=0.5, on_change=lambda: st.session_state.update(bridge_solved=False))
    depth = c3.number_input("Truss Depth (m)", value=3.0, step=0.5, on_change=lambda: st.session_state.update(bridge_solved=False))
    
    sec_list = list(SECTIONS_DB.keys()) if SECTIONS_DB else ["HEA200", "HEB300", "Soldier U100"]
    
    c4, c5, c6, c7 = st.columns(4)
    top_sec = c4.selectbox("Top Chord Section", sec_list, index=0, on_change=lambda: st.session_state.update(bridge_solved=False))
    bot_sec = c5.selectbox("Bottom Chord Section", sec_list, index=0, on_change=lambda: st.session_state.update(bridge_solved=False))
    diag_sec = c6.selectbox("Diagonal Section", sec_list, index=len(sec_list)-1 if len(sec_list)>1 else 0, on_change=lambda: st.session_state.update(bridge_solved=False))
    vert_sec = c7.selectbox("Vertical Section", sec_list, index=len(sec_list)-1 if len(sec_list)>1 else 0, on_change=lambda: st.session_state.update(bridge_solved=False))
    
    st.markdown("#### 🎯 2. Applied Loads")
    w_load = st.number_input("Uniform Load on Bottom Chord (kN/m)", value=25.0, step=1.0, on_change=lambda: st.session_state.update(bridge_solved=False))
    
    st.markdown("---")
    
    if st.button("🚀 Run Bridge FEA & Generate Report", type="primary", use_container_width=True):
        with st.spinner("Building Bridge Mesh & Solving Matrix..."):
            nodes, elements, nodal_loads, L_tot, depth, display_nodes, supports_list = build_bridge_mesh(
                int(num_bays), bay_length, depth, top_sec, bot_sec, diag_sec, vert_sec, w_load
            )
            
            U, R = solve_fea_engine(nodes, elements, nodal_loads, supports_list)
            
            st.session_state.bridge_fea_data = {
                'U': U, 'R': R, 'nodes': nodes, 'elements': elements, 'display_nodes': display_nodes, 'supports_list': supports_list,
                'sys_data': {
                    'L_tot': L_tot, 'depth': depth, 'top_sec': top_sec, 'bot_sec': bot_sec, 
                    'diag_sec': diag_sec, 'vert_sec': vert_sec, 'w_load': w_load
                }
            }
            st.session_state.bridge_solved = True
            
    if st.session_state.bridge_solved:
        fea_data = st.session_state.bridge_fea_data
        
        st.markdown("### 🎛️ Analysis Results & Diagrams (SAP2000 Style)")
        with st.expander("⚙️ Diagram Scale Controls", expanded=True):
            cs1, cs2, cs3 = st.columns(3)
            sc_n = cs1.slider("Axial Scale", 0.001, 0.050, 0.005, step=0.001)
            sc_v = cs2.slider("Shear Scale", 0.001, 0.050, 0.010, step=0.001)
            sc_m = cs3.slider("Moment Scale", 0.001, 0.100, 0.010, step=0.001)
            scales = {'N': sc_n, 'V': sc_v, 'M': sc_m}
            
        img_bufs = plot_bridge_diagrams(
            fea_data['nodes'], fea_data['elements'], fea_data['R'], 
            scales, fea_data['display_nodes'], fea_data['sys_data']['L_tot'], 
            fea_data['supports_list']
        )
        
        fea_data['sys_data']['img_bufs'] = img_bufs
        
        st.image(img_bufs['N'], caption="Axial Force Diagram")
        st.image(img_bufs['M'], caption="Bending Moment Diagram")
        
        docx_out = generate_bridge_report(fea_data['sys_data'])
        st.download_button("⬇️ Download Bridge Calculation Sheet", 
                           data=docx_out.getvalue(), 
                           file_name="Bridge_Structure_Report.docx", 
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
