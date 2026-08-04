# inclined_master.py

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import re
import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from config import SECTIONS_DB, STRUTS_DB
except ImportError:
    st.error("⚠️ برجاء التأكد من وجود ملف config.py")
    SECTIONS_DB = {}
    STRUTS_DB = {}

# =========================================================
# 0. Helper Functions & Styles
# =========================================================
def get_valid_struts(req_len, struts_db):
    valid = []
    for name, props in struts_db.items():
        min_l, max_l = 0.0, 99.0
        if isinstance(props, dict) and 'min' in props and 'max' in props:
            min_l, max_l = props['min'], props['max']
        else:
            m = re.search(r"\((\d+\.?\d*):(\d+\.?\d*)m\)", name)
            if m:
                min_l, max_l = float(m.group(1)), float(m.group(2))
        
        if min_l <= req_len <= max_l:
            valid.append(name)
            
    if not valid:
        return list(struts_db.keys()) if struts_db else ["PPH (Fallback)"]
        
    def priority(name):
        n = name.upper()
        if "PPH" in n: return 1
        if "PPS" in n: return 2
        if "TILT" in n: return 3
        if "MMP" in n: return 4
        return 5
        
    return sorted(valid, key=priority)

def apply_plot_styles():
    mpl.rcParams['font.family'] = 'sans-serif'
    mpl.rcParams['font.sans-serif'] = ['Arial', 'Helvetica', 'DejaVu Sans']
    mpl.rcParams['axes.linewidth'] = 0.3
    mpl.rcParams['font.size'] = 7
    mpl.rcParams['font.weight'] = 'normal'

# =========================================================
# 1. Geometry & Mesh Generator
# =========================================================
def build_fea_mesh(L_segs, L_rem, X_segs, X_rem, angle_rad, applied_loads, inc_sec, base_sec, strut_types, corner_sup, base_sups):
    inc_key_pts = [0.0]
    curr = 0.0
    for seg in L_segs:
        curr += seg
        inc_key_pts.append(curr)
    if L_rem > 0:
        inc_key_pts.append(curr + L_rem)
    
    L_tot = curr + L_rem
    
    for ld in applied_loads:
        inc_key_pts.append(ld['start'])
        if ld['type'] != 'Point Load':
            inc_key_pts.append(ld['end'])
            
    inc_key_pts = sorted(list(set([round(p, 4) for p in inc_key_pts if 0 <= p <= L_tot + 1e-5])))
    
    inc_nodes_L = []
    for i in range(len(inc_key_pts)-1):
        A = inc_key_pts[i]
        B = inc_key_pts[i+1]
        num_sub = max(1, int(np.ceil((B - A) / 0.15))) 
        pts = np.linspace(A, B, num_sub+1)
        for p in pts[:-1]:
            inc_nodes_L.append(p)
    inc_nodes_L.append(inc_key_pts[-1])
    
    nodes = []
    inc_node_indices = []
    
    nodes.append([0.0, 0.0])
    inc_node_indices.append(0)
    
    for L_val in inc_nodes_L[1:]:
        nodes.append([L_val * np.cos(angle_rad), L_val * np.sin(angle_rad)])
        inc_node_indices.append(len(nodes)-1)
        
    X_cum = 0.0
    base_x_pts = [0.0]
    for X_seg in X_segs:
        X_cum += X_seg
        base_x_pts.append(X_cum)
    if X_rem > 0:
        base_x_pts.append(X_cum + X_rem)
        
    for sup in base_sups:
        base_x_pts.append(sup['x'])
        
    base_x_pts = sorted(list(set([round(x, 4) for x in base_x_pts])))
    
    base_node_indices = []
    for x in base_x_pts:
        if x == 0.0:
            base_node_indices.append(0)
        else:
            nodes.append([x, 0.0])
            base_node_indices.append(len(nodes)-1)
            
    supports_list = []
    supports_list.append({'node': 0, 'type': corner_sup['type'], 'angle': corner_sup.get('angle', 0.0)})
    for sup in base_sups:
        idx = base_x_pts.index(round(sup['x'], 4))
        n_idx = base_node_indices[idx]
        supports_list.append({'node': n_idx, 'type': sup['type'], 'angle': sup.get('angle', 0.0)})
        
    display_nodes = set([s['node'] for s in supports_list])
    display_nodes.add(inc_node_indices[-1])
    display_nodes.add(base_node_indices[-1])
    
    target_Ls = [sum(L_segs[:j+1]) for j in range(len(L_segs))]
    for j in range(len(L_segs)):
        target_L = round(target_Ls[j], 4)
        if target_L in inc_nodes_L:
            idx = inc_nodes_L.index(target_L)
            display_nodes.add(inc_node_indices[idx])
            
    for ld in applied_loads:
        if ld['type'] == 'Point Load':
            try:
                idx = inc_nodes_L.index(round(ld['start'], 4))
                display_nodes.add(inc_node_indices[idx])
            except ValueError: pass

    elements = []
    nodal_loads = []
    E_st = 210000000.0 
    inc_props = SECTIONS_DB.get(inc_sec, {'E': E_st, 'A': 0.00343, 'I': 0.00000122})
    
    for i in range(len(inc_node_indices)-1):
        n1 = inc_node_indices[i]
        n2 = inc_node_indices[i+1]
        L_mid = (inc_nodes_L[i] + inc_nodes_L[i+1]) / 2.0
        
        p_x1, p_y1 = 0.0, 0.0
        p_x2, p_y2 = 0.0, 0.0
        
        for ld in applied_loads:
            if ld['type'] == 'Point Load': continue
            if ld['start'] - 1e-4 <= L_mid <= ld['end'] + 1e-4:
                L_len = max(ld['end'] - ld['start'], 1e-5)
                w_a = ld['w1'] + (ld['w2'] - ld['w1']) * (inc_nodes_L[i] - ld['start']) / L_len
                w_b = ld['w1'] + (ld['w2'] - ld['w1']) * (inc_nodes_L[i+1] - ld['start']) / L_len
                
                if ld['dir'] == 'Gravity (Vertical ↓)':
                    c, s = np.cos(angle_rad), np.sin(angle_rad)
                    p_x1 += -w_a * s; p_y1 += -w_a * c
                    p_x2 += -w_b * s; p_y2 += -w_b * c
                else:
                    p_y1 += -w_a; p_y2 += -w_b
                    
        elements.append({
            'type': 'frame', 'group': 'inclined', 'sec': inc_sec,
            'n1': n1, 'n2': n2, 'px1': p_x1, 'py1': p_y1, 'px2': p_x2, 'py2': p_y2,
            'E': inc_props.get('E', E_st), 'A': inc_props.get('A', 0.00343), 'I': inc_props.get('I', 0.00000122)
        })
        
    for ld in applied_loads:
        if ld['type'] == 'Point Load':
            try:
                idx = inc_nodes_L.index(round(ld['start'], 4))
                n_idx = inc_node_indices[idx]
                c, s = np.cos(angle_rad), np.sin(angle_rad)
                if ld['dir'] == 'Gravity (Vertical ↓)':
                    nodal_loads.append({'node': n_idx, 'Fx': 0.0, 'Fy': -ld['w1']})
                else:
                    nodal_loads.append({'node': n_idx, 'Fx': ld['w1']*s, 'Fy': -ld['w1']*c})
            except ValueError: pass
                
    base_props = SECTIONS_DB.get(base_sec, {'E': E_st, 'A': 0.00343, 'I': 0.00000122})
    for i in range(len(base_node_indices)-1):
        elements.append({
            'type': 'frame', 'group': 'base', 'sec': base_sec,
            'n1': base_node_indices[i], 'n2': base_node_indices[i+1],
            'px1': 0.0, 'py1': 0.0, 'px2': 0.0, 'py2': 0.0,
            'E': base_props.get('E', E_st), 'A': base_props.get('A', 0.00343), 'I': base_props.get('I', 0.00000122)
        })
        
    # Connect Struts to base X locations
    X_cum_strut = 0.0
    for j in range(len(L_segs)):
        target_L = round(target_Ls[j], 4)
        X_cum_strut += X_segs[j]
        if target_L in inc_nodes_L:
            idx_inc = inc_nodes_L.index(target_L)
            n_inc = inc_node_indices[idx_inc]
            
            idx_base = base_x_pts.index(round(X_cum_strut, 4))
            n_base = base_node_indices[idx_base]
            display_nodes.add(n_base)
            
            st_props = STRUTS_DB.get(strut_types[j], {'A': 0.001}) 
            elements.append({
                'type': 'truss', 'group': 'strut', 'sec': strut_types[j],
                'n1': n_base, 'n2': n_inc,
                'E': E_st, 'A': st_props.get('A', 0.001)
            })
            
    X_tot = base_x_pts[-1]
    return nodes, elements, nodal_loads, L_tot, X_tot, display_nodes, supports_list

# =========================================================
# 2. Advanced 2D Frame FEA Solver
# =========================================================
def solve_fea_engine(nodes, elements, nodal_loads, supports_list):
    num_nodes = len(nodes)
    NDOF = num_nodes * 3
    K = np.zeros((NDOF, NDOF))
    F = np.zeros(NDOF)
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        x1, y1 = nodes[n1][0], nodes[n1][1]
        x2, y2 = nodes[n2][0], nodes[n2][1]
        L = np.hypot(x2 - x1, y2 - y1)
        if L < 1e-5: continue
        c, s = (x2 - x1) / L, (y2 - y1) / L
        el['L'], el['c'], el['s'] = L, c, s
        
        E_raw = el['E']; E = E_raw if E_raw > 1000000 else E_raw * 10000.0 
        A_raw = el['A']; A = A_raw if A_raw < 0.1 else A_raw / 10000.0 
        
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        
        if el['type'] == 'truss':
            k_loc = np.zeros((6, 6))
            k_loc[0, 0] = k_loc[3, 3] = E * A / L
            k_loc[0, 3] = k_loc[3, 0] = -E * A / L
        else:
            I_raw = el['I']; I = I_raw if I_raw < 0.001 else I_raw / 100000000.0 
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0],
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L],
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2],
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            px1, py1, px2, py2 = el['px1'], el['py1'], el['px2'], el['py2']
            f_eq_loc = np.array([
                (2*px1 + px2)*L/6.0,
                (7*py1 + 3*py2)*L/20.0,
                (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0,
                (3*py1 + 7*py2)*L/20.0,
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_eq_glob = T.T @ f_eq_loc
            dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
            for r in range(6): F[dof_idx[r]] += f_eq_glob[r]
            
        k_glob = T.T @ k_loc @ T
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        for r in range(6):
            for col in range(6):
                K[dof_idx[r], dof_idx[col]] += k_glob[r, col]
                
    for nl in nodal_loads:
        F[3*nl['node'] + 0] += nl['Fx']
        F[3*nl['node'] + 1] += nl['Fy']
            
    K_orig = K.copy()
    fixed_dofs = []
    K_penalty = 1e12
    
    # 💡 تطبيق الركائز المتطورة (Advanced Boundary Conditions)
    for sup in supports_list:
        n = sup['node']
        t = sup['type']
        a = sup['angle']
        if t == 'Fixed':
            fixed_dofs.extend([3*n, 3*n+1, 3*n+2])
        elif t == 'Hinged':
            fixed_dofs.extend([3*n, 3*n+1])
        elif t == 'Roller':
            if abs(a % 180) < 1e-5: 
                fixed_dofs.append(3*n+1)
            elif abs((a - 90) % 180) < 1e-5: 
                fixed_dofs.append(3*n)
            else:
                rad = np.radians(a)
                nx, ny = -np.sin(rad), np.cos(rad) 
                K[3*n, 3*n] += K_penalty * nx**2
                K[3*n+1, 3*n+1] += K_penalty * ny**2
                K[3*n, 3*n+1] += K_penalty * nx * ny
                K[3*n+1, 3*n] += K_penalty * nx * ny

    free_dof = [i for i in range(NDOF) if i not in fixed_dofs]
        
    K_ff = K[np.ix_(free_dof, free_dof)]
    F_f = F[free_dof]
    
    U = np.zeros(NDOF)
    try:
        U_f = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError:
        U_f = np.linalg.lstsq(K_ff, F_f, rcond=None)[0]
    
    U[free_dof] = U_f
    R_reactions = K_orig @ U - F 
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        c, s, L = el['c'], el['s'], el['L']
        E_raw = el['E']; E = E_raw if E_raw > 1000000 else E_raw * 10000.0 
        A_raw = el['A']; A = A_raw if A_raw < 0.1 else A_raw / 10000.0 
        
        dof_idx = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        u_glob = U[dof_idx]
        T = np.array([
            [c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]
        ])
        u_loc = T @ u_glob
        
        if el['type'] == 'truss':
            N_val = (E * A / L) * (u_loc[3] - u_loc[0])
            el['internal'] = {'N': [N_val, N_val], 'V': [0,0], 'M': [0,0], 'x': [0, L]}
        else:
            I_raw = el['I']; I = I_raw if I_raw < 0.001 else I_raw / 100000000.0 
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0],
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L],
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2],
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            px1, py1, px2, py2 = el['px1'], el['py1'], el['px2'], el['py2']
            f_eq_loc = np.array([
                (2*px1 + px2)*L/6.0,
                (7*py1 + 3*py2)*L/20.0,
                (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0,
                (3*py1 + 7*py2)*L/20.0,
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_end = k_loc @ u_loc - f_eq_loc
            
            # 💡 دقة أعلى للحصول على الـ Peaks بشكل مثالي
            xs = np.linspace(0, L, 51) 
            N_arr = np.zeros_like(xs)
            V_arr = np.zeros_like(xs)
            M_arr = np.zeros_like(xs)
            
            for i, x in enumerate(xs):
                N_arr[i] = -f_end[0] - (px1*x + (px2-px1)*x**2/(2*L))
                V_arr[i] = f_end[1] + (py1*x + (py2-py1)*x**2/(2*L))
                M_arr[i] = -f_end[2] + f_end[1]*x + py1*x**2/2.0 + (py2-py1)*x**3/(6*L)
                
            el['internal'] = {'N': N_arr, 'V': V_arr, 'M': M_arr, 'x': xs}
            
    return U, R_reactions

# =========================================================
# 3. Plotting Engines (Independent Subplots)
# =========================================================
def get_img_buf(fig):
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', transparent=True)
    plt.close(fig)
    buf.seek(0)
    return buf

def draw_base_geometry(ax, nodes, elements, supports_list):
    # رسم خطوط العناصر
    for el in elements:
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        color = 'black' if el['type'] == 'frame' else 'gray'
        style = '-' if el['type'] == 'frame' else '--'
        ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color=color, linestyle=style, linewidth=0.5, zorder=1)
        
    # رسم الركائز بأشكال هندسية دقيقة
    for sup in supports_list:
        n = sup['node']
        x, y = nodes[n][0], nodes[n][1]
        t = sup['type']
        a = sup['angle']
        if t == 'Fixed':
            ax.plot(x, y, marker='s', color='blue', markersize=6, zorder=5)
        elif t == 'Hinged':
            ax.plot(x, y, marker='^', color='orange', markersize=8, zorder=5)
        elif t == 'Roller':
            ax.plot(x, y, marker='o', color='green', markersize=6, zorder=5)
            rad = np.radians(a)
            dx, dy = np.cos(rad), np.sin(rad)
            ax.plot([x - 0.4*dx, x + 0.4*dx], [y - 0.4*dy, y + 0.4*dy], color='green', lw=1.5, zorder=4)

def draw_section_names(ax, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec, is_n_diagram=False):
    angle_rad = np.radians(angle_deg)
    c_ang, s_ang = np.cos(angle_rad), np.sin(angle_rad)
    
    inc_mid_x = (L_tot/2) * c_ang
    inc_mid_y = (L_tot/2) * s_ang
    ax.text(inc_mid_x - s_ang*0.6, inc_mid_y + c_ang*0.6, inc_sec, color='gray', fontsize=7, alpha=0.9, ha='center', va='center', rotation=angle_deg, fontname='Arial')
    
    base_mid_x = X_tot/2
    ax.text(base_mid_x, -0.6, base_sec, color='gray', fontsize=7, alpha=0.9, ha='center', va='center', fontname='Arial')
    
    drawn_struts = set()
    for el in elements:
        if el['group'] == 'strut':
            sig = f"{el['n1']}_{el['n2']}"
            if sig not in drawn_struts:
                n1, n2 = nodes[el['n1']], nodes[el['n2']]
                dx, dy = n2[0]-n1[0], n2[1]-n1[1]
                L_s = np.hypot(dx, dy)
                nx, ny = -dy/L_s, dx/L_s
                mid_x, mid_y = (n1[0]+n2[0])/2, (n1[1]+n2[1])/2
                rot = np.degrees(np.arctan2(dy, dx))
                
                if is_n_diagram:
                    ax.text(mid_x - nx*0.4, mid_y - ny*0.4, el['sec'], color='gray', fontsize=6, alpha=0.9, ha='center', va='center', rotation=rot, fontname='Arial')
                else:
                    ax.text(mid_x + nx*0.2, mid_y + ny*0.2, el['sec'], color='gray', fontsize=6, alpha=0.9, ha='center', va='center', rotation=rot, fontname='Arial')
                drawn_struts.add(sig)

def plot_live_geometry(nodes, elements, applied_loads, L_segs, X_segs, angle_deg, inc_sec, base_sec, L_tot, X_tot, supports_list):
    apply_plot_styles()
    fig, ax = plt.subplots(figsize=(6, 5))
    ax.set_aspect('equal', adjustable='datalim')
    ax.axis('off')
    
    draw_base_geometry(ax, nodes, elements, supports_list)
    draw_section_names(ax, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec)

    angle_rad = np.radians(angle_deg)
    c_ang, s_ang = np.cos(angle_rad), np.sin(angle_rad)
    
    curr_l = 0.0
    for i, seg in enumerate(L_segs):
        px = (curr_l + seg/2) * c_ang
        py = (curr_l + seg/2) * s_ang
        ax.text(px - s_ang*0.9, py + c_ang*0.9, f"L{i+1}={seg:.2f}m", color='gray', fontsize=7, rotation=angle_deg, ha='center', va='center', fontname='Arial')
        curr_l += seg
        
    curr_x = 0.0
    for i, seg in enumerate(X_segs):
        px = curr_x + seg/2
        py = 0.0
        ax.text(px, py - 0.9, f"X{i+1}={seg:.2f}m", color='gray', fontsize=7, ha='center', va='center', fontname='Arial')
        curr_x += seg

    if applied_loads:
        max_w = max([max(abs(ld['w1']), abs(ld['w2'])) for ld in applied_loads] + [1.0])
        scale_ld = 1.2 / max_w
        for ld in applied_loads:
            w1, w2 = ld['w1'], ld['w2']
            start_L, end_L = ld['start'], ld['end']
            dir_type = ld['dir']
            px1, py1 = start_L * c_ang, start_L * s_ang
            px2, py2 = end_L * c_ang, end_L * s_ang
            
            if ld['type'] == 'Point Load':
                arrow_len = 1.0
                if dir_type == 'Gravity (Vertical ↓)':
                    ax.arrow(px1, py1 + arrow_len + 0.1, 0, -arrow_len, head_width=0.15, head_length=0.2, length_includes_head=True, fc='fuchsia', ec='fuchsia', zorder=4, linewidth=0.5)
                    ax.text(px1, py1 + arrow_len + 0.3, f"{w1}", color='fuchsia', fontsize=8, ha='center', fontname='Arial')
                else:
                    start_x = px1 - s_ang*(arrow_len+0.1)
                    start_y = py1 + c_ang*(arrow_len+0.1)
                    ax.arrow(start_x, start_y, s_ang*arrow_len, -c_ang*arrow_len, head_width=0.15, head_length=0.2, length_includes_head=True, fc='fuchsia', ec='fuchsia', zorder=4, linewidth=0.5)
                    ax.text(start_x - s_ang*0.2, start_y + c_ang*0.2, f"{w1}", color='fuchsia', fontsize=8, ha='center', rotation=angle_deg, fontname='Arial')
            else:
                if dir_type == 'Gravity (Vertical ↓)':
                    hx1, hy1 = px1, py1 + w1 * scale_ld
                    hx2, hy2 = px2, py2 + w2 * scale_ld
                else:
                    hx1, hy1 = px1 - s_ang * w1 * scale_ld, py1 + c_ang * w1 * scale_ld
                    hx2, hy2 = px2 - s_ang * w2 * scale_ld, py2 + c_ang * w2 * scale_ld
                    
                poly = Polygon([(px1,py1), (hx1,hy1), (hx2,hy2), (px2,py2)], facecolor='none', edgecolor='magenta', linewidth=0.8, zorder=3)
                ax.add_patch(poly)
                
                num_arrows = max(3, int((end_L - start_L) / 0.4))
                xs = np.linspace(start_L, end_L, num_arrows)
                for x_dist in xs:
                    w_curr = w1 + (w2 - w1) * (x_dist - start_L) / max(end_L - start_L, 1e-5)
                    if abs(w_curr) < 0.1: continue
                    px = x_dist * c_ang
                    py = x_dist * s_ang
                    hl = w_curr * scale_ld
                    if dir_type == 'Gravity (Vertical ↓)':
                        ax.arrow(px, py + hl, 0, -hl, head_width=0.08, head_length=0.1, length_includes_head=True, fc='magenta', ec='magenta', linewidth=0.3, zorder=2)
                    else:
                        ax.arrow(px - s_ang*hl, py + c_ang*hl, s_ang*hl, -c_ang*hl, head_width=0.08, head_length=0.1, length_includes_head=True, fc='magenta', ec='magenta', linewidth=0.3, zorder=2)
                        
                if dir_type == 'Gravity (Vertical ↓)':
                    ax.text(px1, py1 + w1*scale_ld + 0.15, f"{w1}", color='magenta', fontsize=7, ha='center', fontname='Arial')
                    ax.text(px2, py2 + w2*scale_ld + 0.15, f"{w2}", color='magenta', fontsize=7, ha='center', fontname='Arial')
                else:
                    ax.text(hx1 - s_ang*0.15, hy1 + c_ang*0.15, f"{w1}", color='magenta', fontsize=7, ha='center', rotation=angle_deg, fontname='Arial')
                    ax.text(hx2 - s_ang*0.15, hy2 + c_ang*0.15, f"{w2}", color='magenta', fontsize=7, ha='center', rotation=angle_deg, fontname='Arial')

    return get_img_buf(fig)

def plot_sap2000_diagrams(nodes, elements, R_reactions, scales, display_nodes, applied_loads, angle_deg, L_tot, X_tot, inc_sec, base_sec, supports_list):
    apply_plot_styles()
    angle_rad = np.radians(angle_deg)
    c_ang, s_ang = np.cos(angle_rad), np.sin(angle_rad)
    
    figs_dict = {}
    
    # --- 1. Load Diagram ---
    fig_ld, ax_ld = plt.subplots(figsize=(6, 5))
    ax_ld.set_aspect('equal', adjustable='datalim')
    ax_ld.axis('off')
    draw_base_geometry(ax_ld, nodes, elements, supports_list)
    draw_section_names(ax_ld, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec)
    
    if applied_loads:
        max_w = max([max(abs(ld['w1']), abs(ld['w2'])) for ld in applied_loads] + [1.0])
        scale_ld = 1.2 / max_w
        for ld in applied_loads:
            w1, w2 = ld['w1'], ld['w2']
            start_L, end_L = ld['start'], ld['end']
            dir_type = ld['dir']
            px1, py1 = start_L * c_ang, start_L * s_ang
            px2, py2 = end_L * c_ang, end_L * s_ang
            
            if ld['type'] == 'Point Load':
                arrow_len = 1.0
                if dir_type == 'Gravity (Vertical ↓)':
                    ax_ld.arrow(px1, py1 + arrow_len + 0.1, 0, -arrow_len, head_width=0.15, head_length=0.2, length_includes_head=True, fc='fuchsia', ec='fuchsia', zorder=4, linewidth=0.5)
                    ax_ld.text(px1, py1 + arrow_len + 0.3, f"{w1}", color='fuchsia', fontsize=8, ha='center', fontname='Arial')
                else:
                    start_x = px1 - s_ang*(arrow_len+0.1)
                    start_y = py1 + c_ang*(arrow_len+0.1)
                    ax_ld.arrow(start_x, start_y, s_ang*arrow_len, -c_ang*arrow_len, head_width=0.15, head_length=0.2, length_includes_head=True, fc='fuchsia', ec='fuchsia', zorder=4, linewidth=0.5)
                    ax_ld.text(start_x - s_ang*0.2, start_y + c_ang*0.2, f"{w1}", color='fuchsia', fontsize=8, ha='center', rotation=angle_deg, fontname='Arial')
            else:
                if dir_type == 'Gravity (Vertical ↓)':
                    hx1, hy1 = px1, py1 + w1 * scale_ld
                    hx2, hy2 = px2, py2 + w2 * scale_ld
                else:
                    hx1, hy1 = px1 - s_ang * w1 * scale_ld, py1 + c_ang * w1 * scale_ld
                    hx2, hy2 = px2 - s_ang * w2 * scale_ld, py2 + c_ang * w2 * scale_ld
                    
                poly = Polygon([(px1,py1), (hx1,hy1), (hx2,hy2), (px2,py2)], facecolor='none', edgecolor='magenta', linewidth=0.8, zorder=3)
                ax_ld.add_patch(poly)
                
                num_arrows = max(3, int((end_L - start_L) / 0.4))
                xs = np.linspace(start_L, end_L, num_arrows)
                for x_dist in xs:
                    w_curr = w1 + (w2 - w1) * (x_dist - start_L) / max(end_L - start_L, 1e-5)
                    if abs(w_curr) < 0.1: continue
                    px = x_dist * c_ang
                    py = x_dist * s_ang
                    hl = w_curr * scale_ld
                    if dir_type == 'Gravity (Vertical ↓)':
                        ax_ld.arrow(px, py + hl, 0, -hl, head_width=0.08, head_length=0.1, length_includes_head=True, fc='magenta', ec='magenta', linewidth=0.3, zorder=2)
                    else:
                        ax_ld.arrow(px - s_ang*hl, py + c_ang*hl, s_ang*hl, -c_ang*hl, head_width=0.08, head_length=0.1, length_includes_head=True, fc='magenta', ec='magenta', linewidth=0.3, zorder=2)
                        
                if dir_type == 'Gravity (Vertical ↓)':
                    ax_ld.text(px1, py1 + w1*scale_ld + 0.15, f"{w1}", color='magenta', fontsize=7, ha='center', fontname='Arial')
                    ax_ld.text(px2, py2 + w2*scale_ld + 0.15, f"{w2}", color='magenta', fontsize=7, ha='center', fontname='Arial')
                else:
                    ax_ld.text(hx1 - s_ang*0.15, hy1 + c_ang*0.15, f"{w1}", color='magenta', fontsize=7, ha='center', rotation=angle_deg, fontname='Arial')
                    ax_ld.text(hx2 - s_ang*0.15, hy2 + c_ang*0.15, f"{w2}", color='magenta', fontsize=7, ha='center', rotation=angle_deg, fontname='Arial')
    figs_dict['Load'] = get_img_buf(fig_ld)

    # --- 2. Reactions Diagram ---
    fig_react, ax_react = plt.subplots(figsize=(6, 5))
    ax_react.set_aspect('equal', adjustable='datalim')
    ax_react.axis('off')
    draw_base_geometry(ax_react, nodes, elements, supports_list)
    draw_section_names(ax_react, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec)
    
    for i, n in enumerate(nodes):
        Rx, Ry = R_reactions[3*i], R_reactions[3*i+1]
        x, y = n[0], n[1]
        if abs(Rx) > 0.1:
            ax_react.annotate("", xy=(x, y), xytext=(x - 0.7*np.sign(Rx), y), arrowprops=dict(color='purple', width=0.5, headwidth=4, headlength=4))
            ax_react.text(x - 0.8*np.sign(Rx), y - 0.2, f"{abs(Rx):.1f}", color='purple', fontname='Arial', fontsize=8, ha='center', va='top')
        if abs(Ry) > 0.1:
            ax_react.annotate("", xy=(x, y), xytext=(x, y - 0.7*np.sign(Ry)), arrowprops=dict(color='purple', width=0.5, headwidth=4, headlength=4))
            ax_react.text(x, y - 0.8*np.sign(Ry), f"{abs(Ry):.1f}", color='purple', fontname='Arial', fontsize=8, ha='center', va='top')
    figs_dict['React'] = get_img_buf(fig_react)

    # --- Force Diagrams Helper ---
    def create_force_diagram(val_key, scale, color_pos, color_neg):
        fig_f, ax_f = plt.subplots(figsize=(6, 5))
        ax_f.set_aspect('equal', adjustable='datalim')
        ax_f.axis('off')
        
        draw_base_geometry(ax_f, nodes, elements, supports_list)
        draw_section_names(ax_f, elements, nodes, L_tot, X_tot, angle_deg, inc_sec, base_sec, is_n_diagram=(val_key=='N'))
        
        plotted_texts = set()
        def write_val(txt_x, txt_y, v, rot=0):
            if abs(v) >= 0.1:
                lbl = f"{abs(v):.1f}"
                sig = f"{round(txt_x,1)}_{round(txt_y,1)}"
                if sig not in plotted_texts:
                    ax_f.text(txt_x, txt_y, lbl, color='black', fontsize=7, fontname='Arial', ha='center', va='center', rotation=rot)
                    plotted_texts.add(sig)

        for el in elements:
            n1, n2 = nodes[el['n1']], nodes[el['n2']]
            x1, y1 = n1[0], n1[1]
            x2, y2 = n2[0], n2[1]
            dx, dy = x2 - x1, y2 - y1
            L_s = np.hypot(dx, dy)
            if L_s < 1e-5: continue
            
            c, s = dx/L_s, dy/L_s
            rot_ang = np.degrees(np.arctan2(dy, dx))
            
            if el['type'] == 'truss' and val_key == 'N':
                val = el['internal']['N'][0]
                if abs(val) < 0.1: continue
                nx, ny = -dy/L_s, dx/L_s
                h = max(0.4, abs(val) * scale) 
                color = color_pos if val >= 0 else color_neg
                
                p1, p2 = (x1, y1), (x2, y2)
                p3, p4 = (x2 + nx * h, y2 + ny * h), (x1 + nx * h, y1 + ny * h)
                
                ax_f.add_patch(Polygon([p1, p2, p3, p4], facecolor='none', edgecolor=color, linewidth=0.8, zorder=2))
                
                num_lines = max(5, int(L_s / 0.3))
                for i in range(1, num_lines):
                    frac = i / num_lines
                    lx, ly = x1 + frac * dx, y1 + frac * dy
                    ax_f.plot([lx, lx + nx * h], [ly, ly + ny * h], color=color, linewidth=0.3, alpha=0.6)
                    
                mid_h_x, mid_h_y = x1 + dx/2 + nx*h/2, y1 + dy/2 + ny*h/2
                write_val(mid_h_x, mid_h_y, val, rot_ang)
                continue
                
            if el['type'] == 'frame':
                xs_arr = el['internal']['x']
                vals_orig = el['internal'][val_key]
                # 💡 عكس اتجاه العزوم لتُرسم ناحية الشد
                plot_vals = -vals_orig if val_key == 'M' else vals_orig
                
                px_arr = x1 + c * xs_arr - s * plot_vals * scale
                py_arr = y1 + s * xs_arr + c * plot_vals * scale
                
                ax_f.plot(np.append(x1, np.append(px_arr, x2)), np.append(y1, np.append(py_arr, y2)), color=color_pos, linewidth=0.8)
                
                num_lines = max(2, int(L_s / 0.4))
                for i in range(1, num_lines):
                    frac = i / num_lines
                    lx, ly = x1 + frac * dx, y1 + frac * dy
                    idx_val = int(frac * (len(plot_vals)-1))
                    lv = plot_vals[idx_val]
                    hx, hy = lx - s * lv * scale, ly + c * lv * scale
                    line_color = color_pos if lv >= 0 else color_neg
                    ax_f.plot([lx, hx], [ly, hy], color=line_color, linewidth=0.3, alpha=0.6)
                    
                offset = 0.25
                v_start = plot_vals[0]
                if el['n1'] in display_nodes:
                    txt_x = x1 - s * v_start * scale - s * np.sign(v_start) * offset
                    txt_y = y1 + c * v_start * scale + c * np.sign(v_start) * offset
                    write_val(txt_x, txt_y, vals_orig[0])
                    
                v_end = plot_vals[-1]
                if el['n2'] in display_nodes:
                    txt_x = x2 - s * v_end * scale - s * np.sign(v_end) * offset
                    txt_y = y2 + c * v_end * scale + c * np.sign(v_end) * offset
                    write_val(txt_x, txt_y, vals_orig[-1])
                    
                # 💡 اكتشاف ورسم قمم العزوم الموجبة والسالبة (Mid-span Peaks)
                if val_key == 'M':
                    for i in range(1, len(plot_vals)-1):
                        v_prev, v_curr, v_next = plot_vals[i-1], plot_vals[i], plot_vals[i+1]
                        if (v_curr > v_prev and v_curr > v_next) or (v_curr < v_prev and v_curr < v_next):
                            if abs(v_curr) > 0.1 and abs(v_curr) > 0.05 * max(abs(plot_vals)):
                                txt_x = x1 + c * xs_arr[i] - s * v_curr * scale - s * np.sign(v_curr) * offset
                                txt_y = y1 + s * xs_arr[i] + c * v_curr * scale + c * np.sign(v_curr) * offset
                                write_val(txt_x, txt_y, vals_orig[i])
                                
        return get_img_buf(fig_f)

    # --- 3. Generate N, V, M ---
    figs_dict['N'] = create_force_diagram('N', scales['N'], 'blue', 'red')
    figs_dict['V'] = create_force_diagram('V', scales['V'], 'purple', 'magenta')
    figs_dict['M'] = create_force_diagram('M', scales['M'], 'green', 'y')

    return figs_dict

# =========================================================
# 4. Report Generator for Inclined Systems
# =========================================================
def generate_inclined_report(sys_data):
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
        doc.add_page_break()
    else:
        doc = Document()
        
    def force_ltr_left(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)
        
    def add_line(text, bold=False):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.rtl = False
        
    def add_check(component, param, act, allw, unit):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.5
        r_title = p.add_run(f"• Check {component} ({param}):\n")
        r_title.bold = True
        r_title.font.rtl = False
        r_act = p.add_run(f"  Actual = {act:.2f} {unit}  <  Allowable = {allw:.2f} {unit}  ")
        r_act.font.rtl = False
        res = "SAFE" if act <= allw else "UNSAFE"
        r_res = p.add_run(res)
        r_res.font.bold = True
        r_res.font.rtl = False
        r_res.font.color.rgb = RGBColor(0, 128, 0) if res == "SAFE" else RGBColor(255, 0, 0)
    
    p_title = doc.add_paragraph()
    force_ltr_left(p_title)
    run_title = p_title.add_run("CALCULATION SHEET FOR INCLINED FORMWORK SYSTEM")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.rtl = False
    
    add_line("="*50, bold=True)
    
    add_line(f"1. Geometry & Inputs:", bold=True)
    add_line(f"- Inclined Soldier Length = {sys_data['L_tot']:.2f} m")
    add_line(f"- Inclination Angle = {sys_data['angle']:.1f} degrees")
    add_line(f"- Applied Loads = Variable (Refer to Diagram)")
    
    doc.add_paragraph()
    add_line(f"2. Safety Checks:", bold=True)
    
    inc_sec = sys_data['inc_sec']
    add_line(f"A. Inclined Soldier ({inc_sec})", bold=True)
    add_check("Moment", "M_max", sys_data['max_M_inc'], SECTIONS_DB.get(inc_sec, {}).get('Mall', 999), "kN.m")
    add_check("Shear", "V_max", sys_data['max_V_inc'], SECTIONS_DB.get(inc_sec, {}).get('Qall', 999), "kN")
    
    base_sec = sys_data['base_sec']
    add_line(f"B. Horizontal Base Soldier ({base_sec})", bold=True)
    add_check("Moment", "M_max", sys_data['max_M_base'], SECTIONS_DB.get(base_sec, {}).get('Mall', 999), "kN.m")
    add_check("Shear", "V_max", sys_data['max_V_base'], SECTIONS_DB.get(base_sec, {}).get('Qall', 999), "kN")
    
    add_line("C. Push-Pull Struts (Axial Force)", bold=True)
    for idx, st_val in enumerate(sys_data['struts_res']):
        allow = STRUTS_DB.get(st_val['type'], {}).get('allow', 999) 
        add_check(f"Strut {idx+1} ({st_val['type']})", "N_max", st_val['N'], allow, "kN")
        
    doc.add_page_break()
    add_line("3. Analysis Diagrams:", bold=True)
    
    titles = {
        'Load': "Assigned Load Diagram",
        'React': "Reactions Diagram (kN)",
        'N': "Axial Force Diagram (kN)",
        'V': "Shear Force Diagram (kN)",
        'M': "Bending Moment Diagram (kN.m)"
    }
    
    for key in ['Load', 'React', 'N', 'V', 'M']:
        buf = sys_data['img_bufs'][key]
        buf.seek(0)
        
        p_img = doc.add_paragraph()
        force_ltr_left(p_img)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(buf.read()), width=Cm(14.0))
        
        p_txt = doc.add_paragraph()
        force_ltr_left(p_txt)
        p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_txt = p_txt.add_run(titles[key])
        r_txt.font.name = 'Arial'
        r_txt.font.size = Pt(11)
        r_txt.underline = True
        r_txt.font.rtl = False
        
        doc.add_page_break()
    
    out = io.BytesIO()
    doc.save(out)
    return out

# =========================================================
# 5. Main UI Module for Inclined Systems
# =========================================================
def render_inclined_module():
    st.markdown("## 📐 Inclined Formwork System (Advanced FEA)")
    
    if 'inclined_solved' not in st.session_state:
        st.session_state.inclined_solved = False
        
    c_top1, c_top2 = st.columns(2)
    angle_deg = c_top1.number_input("Inclination Angle (Degrees, < 90)", value=60.0, step=5.0, on_change=lambda: st.session_state.update(inclined_solved=False))
    angle_rad = np.radians(angle_deg)
    num_struts = c_top2.number_input("Number of Push-Pulls", min_value=1, max_value=5, value=2, step=1, on_change=lambda: st.session_state.update(inclined_solved=False))
    
    st.markdown("---")
    
    c_p1, c_p2 = st.columns(2)
    sec_list = list(SECTIONS_DB.keys()) if SECTIONS_DB else ["Soldier U100"]
    default_idx = next((i for i, sec in enumerate(sec_list) if 'Soldier' in sec), 0)
    
    inc_sec = c_p1.selectbox("Profile (Inclined)", sec_list, index=default_idx, on_change=lambda: st.session_state.update(inclined_solved=False))
    base_sec = c_p2.selectbox("Profile (Base)", sec_list, index=default_idx, on_change=lambda: st.session_state.update(inclined_solved=False))
    
    st.markdown("#### ⚓ 1. Geometry & Struts Segments")
    L_segs, X_segs, strut_types = [], [], []
    L_cum, X_cum = 0.0, 0.0
    
    for j in range(int(num_struts)):
        cl1, cl2, cl3 = st.columns([1, 1, 1.5])
        l_val = cl1.number_input(f"L{j+1} on Inclined (m)", value=2.0, step=0.5, key=f"L_{j}", on_change=lambda: st.session_state.update(inclined_solved=False))
        x_val = cl2.number_input(f"X{j+1} on Base (m)", value=1.5, step=0.5, key=f"X_{j}", on_change=lambda: st.session_state.update(inclined_solved=False))
        
        L_segs.append(l_val)
        X_segs.append(x_val)
        L_cum += l_val
        X_cum += x_val
        
        req_len = np.hypot(X_cum - L_cum * np.cos(angle_rad), 0 - L_cum * np.sin(angle_rad))
        valid_struts = get_valid_struts(req_len, STRUTS_DB)
        st_type = cl3.selectbox(f"Strut {j+1} (Req: {req_len:.2f}m)", valid_struts, key=f"st_{j}", on_change=lambda: st.session_state.update(inclined_solved=False))
        strut_types.append(st_type)
        
    cr1, cr2 = st.columns(2)
    L_rem = cr1.number_input("Remaining Inclined Top L (m)", value=1.0, step=0.5, on_change=lambda: st.session_state.update(inclined_solved=False))
    X_rem = cr2.number_input("Remaining Base Right X (m)", value=0.5, step=0.5, on_change=lambda: st.session_state.update(inclined_solved=False))
    
    st.markdown("#### ⚓ 2. Ground Supports Configuration")
    st.info("The corner support is at X=0. You can define additional ground supports anywhere on the base.")
    
    c_sup1, c_sup2 = st.columns(2)
    corner_type = c_sup1.selectbox("Corner Support (Bottom-Left)", ["Hinged", "Roller", "Fixed"], key="corn_type", on_change=lambda: st.session_state.update(inclined_solved=False))
    corner_ang = c_sup2.number_input("Corner Roller Angle (0=Horiz)", value=0.0, step=15.0, key="corn_ang", on_change=lambda: st.session_state.update(inclined_solved=False)) if corner_type == "Roller" else 0.0
    corner_sup = {'type': corner_type, 'angle': corner_ang}
    
    num_base_sups = st.number_input("Number of Additional Ground Supports", 0, 10, int(num_struts), on_change=lambda: st.session_state.update(inclined_solved=False))
    base_sups = []
    
    default_xs = [sum(X_segs[:i+1]) for i in range(len(X_segs))]
    for i in range(int(num_base_sups)):
        cs1, cs2, cs3 = st.columns(3)
        def_x = default_xs[i] if i < len(default_xs) else float((i+1)*1.5)
        sx = cs1.number_input(f"Support {i+1} X (m)", value=def_x, step=0.5, key=f"sx_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
        stype = cs2.selectbox(f"Type {i+1}", ["Hinged", "Roller", "Fixed"], key=f"stype_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
        sang = cs3.number_input(f"Angle {i+1}", value=0.0, step=15.0, key=f"sang_{i}", on_change=lambda: st.session_state.update(inclined_solved=False)) if stype == "Roller" else 0.0
        base_sups.append({'x': sx, 'type': stype, 'angle': sang})

    # 💡 تقسيم الشاشة إلى جزئين (إدخال الأحمال يميناً والشاشة التفاعلية يساراً) لعدم تشتيت العين
    c_in, c_plot = st.columns([1.3, 1])
    
    with c_in:
        st.markdown("#### 🎯 3. Applied Loads on Inclined Soldier")
        num_loads = st.number_input("Number of Load Blocks", 1, 5, 1, on_change=lambda: st.session_state.update(inclined_solved=False))
        applied_loads = []
        for i in range(int(num_loads)):
            with st.expander(f"Load Block {i+1}", expanded=True):
                l_type = st.selectbox("Load Type", ["Uniform", "Trapezoidal/Triangular", "Point Load"], key=f"lt_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
                
                if l_type == "Point Load":
                    c_pt_top1, c_pt_top2 = st.columns(2)
                    num_pts = c_pt_top1.number_input("Number of Point Loads", 1, 20, 1, key=f"npts_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
                    ldir = c_pt_top2.selectbox("Direction", ["Gravity (Vertical ↓)", "Perpendicular (Local ↘)"], key=f"ldir_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
                    
                    st.markdown("<span style='font-size:13px; color:gray;'>Specify distance and value for each point load:</span>", unsafe_allow_html=True)
                    for pt in range(int(num_pts)):
                        c_pt1, c_pt2 = st.columns(2)
                        start_l = c_pt1.number_input(f"Distance {pt+1} from bottom (m)", value=0.0, step=0.5, key=f"ls_{i}_{pt}", on_change=lambda: st.session_state.update(inclined_solved=False))
                        w1 = c_pt2.number_input(f"Load {pt+1} Value (kN)", value=15.0, step=1.0, key=f"w1_{i}_{pt}", on_change=lambda: st.session_state.update(inclined_solved=False))
                        applied_loads.append({'type': l_type, 'start': start_l, 'end': start_l, 'w1': w1, 'w2': w1, 'dir': ldir})
                        
                else:
                    c_top1, c_top2 = st.columns(2)
                    num_items = c_top1.number_input(f"Number of {l_type.split()[0]} Loads", 1, 20, 1, key=f"nitems_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
                    ldir = c_top2.selectbox("Direction", ["Gravity (Vertical ↓)", "Perpendicular (Local ↘)"], key=f"ldir_{i}", on_change=lambda: st.session_state.update(inclined_solved=False))
                    
                    st.markdown("<span style='font-size:13px; color:gray;'>Specify parameters for each load:</span>", unsafe_allow_html=True)
                    for item in range(int(num_items)):
                        if l_type == "Uniform":
                            cl1, cl2, cw1 = st.columns(3)
                            start_l = cl1.number_input(f"Start {item+1} (m)", value=0.0, step=0.5, key=f"ls_{i}_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                            len_l = cl2.number_input(f"Length {item+1} (m)", value=sum(L_segs)+L_rem, step=0.5, key=f"ll_{i}_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                            w1 = cw1.number_input(f"W {item+1} (kN/m)", value=15.0, step=1.0, key=f"w1_{i}_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                            applied_loads.append({'type': l_type, 'start': start_l, 'end': start_l+len_l, 'w1': w1, 'w2': w1, 'dir': ldir})
                        else:
                            cl1, cl2, cw1, cw2 = st.columns(4)
                            start_l = cl1.number_input(f"Start {item+1} (m)", value=0.0, step=0.5, key=f"ls_{i}_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                            len_l = cl2.number_input(f"Length {item+1} (m)", value=sum(L_segs)+L_rem, step=0.5, key=f"ll_{i}_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                            w1 = cw1.number_input(f"W1 {item+1} (kN/m)", value=15.0, step=1.0, key=f"w1_{i}_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                            w2 = cw2.number_input(f"W2 {item+1} (kN/m)", value=0.0, step=1.0, key=f"w2_{i}_{item}", on_change=lambda: st.session_state.update(inclined_solved=False))
                            applied_loads.append({'type': l_type, 'start': start_l, 'end': start_l+len_l, 'w1': w1, 'w2': w2, 'dir': ldir})

    nodes, elements, nodal_loads, L_tot, X_tot, display_nodes, supports_list = build_fea_mesh(L_segs, L_rem, X_segs, X_rem, angle_rad, applied_loads, inc_sec, base_sec, strut_types, corner_sup, base_sups)

    with c_plot:
        st.markdown("<h4 style='text-align: center; font-family: Arial; font-weight: normal; border-bottom: 1px solid gray; padding-bottom: 5px;'>Live Assigned Loads</h4>", unsafe_allow_html=True)
        live_img_buf = plot_live_geometry(nodes, elements, applied_loads, L_segs, X_segs, angle_deg, inc_sec, base_sec, L_tot, X_tot, supports_list)
        st.image(live_img_buf, use_container_width=True)

    st.markdown("---")
    
    col_btn, col_blank = st.columns([1, 2])
    with col_btn:
        if st.button("🚀 Run Advanced FEA & Generate Report", type="primary", use_container_width=True):
            with st.spinner("Building Matrix & Solving FEA..."):
                U, R = solve_fea_engine(nodes, elements, nodal_loads, supports_list)
                st.session_state.inclined_fea_data = {
                    'U': U, 'R': R, 'nodes': nodes, 'elements': elements, 'display_nodes': display_nodes, 'supports_list': supports_list,
                    'sys_data': {
                        'L_tot': L_tot, 'X_tot': X_tot, 'angle': angle_deg, 'W': "Variable", 'ld_dir': "Variable",
                        'inc_sec': inc_sec, 'base_sec': base_sec
                    }
                }
                st.session_state.inclined_solved = True
    
    if st.session_state.inclined_solved:
        fea_data = st.session_state.inclined_fea_data
        
        st.markdown("### 🎛️ Analysis Results & Diagrams")
        with st.expander("⚙️ Diagram Scale Controls", expanded=True):
            c_s1, c_s2, c_s3 = st.columns(3)
            sc_n = c_s1.slider("Axial Force Scale", 0.001, 0.100, 0.02, step=0.005)
            sc_v = c_s2.slider("Shear Force Scale", 0.001, 0.100, 0.02, step=0.005)
            sc_m = c_s3.slider("Moment Scale", 0.01, 0.50, 0.10, step=0.01)
            scales = {'N': sc_n, 'V': sc_v, 'M': sc_m}
            
        img_bufs = plot_sap2000_diagrams(fea_data['nodes'], fea_data['elements'], fea_data['R'], scales, fea_data['display_nodes'], applied_loads, angle_deg, fea_data['sys_data']['L_tot'], fea_data['sys_data']['X_tot'], fea_data['sys_data']['inc_sec'], fea_data['sys_data']['base_sec'], fea_data['supports_list'])
        
        titles = {
            'Load': "Assigned Load Diagram",
            'React': "Reactions Diagram (kN)",
            'N': "Axial Force Diagram (kN)",
            'V': "Shear Force Diagram (kN)",
            'M': "Bending Moment Diagram (kN.m)"
        }
        
        c_p1, c_p2, c_p3 = st.columns(3)
        cols = [c_p1, c_p2, c_p3, c_p1, c_p2]
        
        for idx, key in enumerate(['Load', 'React', 'N', 'V', 'M']):
            with cols[idx]:
                st.image(img_bufs[key], use_container_width=True)
                st.markdown(f"<p style='text-align: center; border-bottom: 1px solid gray; padding-bottom: 5px; font-family: Arial; font-size: 14px;'>{titles[key]}</p>", unsafe_allow_html=True)
        
        max_M_inc, max_V_inc = 0, 0
        max_M_base, max_V_base = 0, 0
        
        for el in fea_data['elements']:
            if el['type'] == 'frame':
                max_M = max(abs(el['internal']['M'][0]), abs(el['internal']['M'][-1]))
                if len(el['internal']['M']) > 2: max_M = max(max_M, np.max(np.abs(el['internal']['M'])))
                max_V = max(abs(el['internal']['V'][0]), abs(el['internal']['V'][-1]))
                if len(el['internal']['V']) > 2: max_V = max(max_V, np.max(np.abs(el['internal']['V'])))
                
                if el['group'] == 'inclined':
                    max_M_inc = max(max_M_inc, max_M); max_V_inc = max(max_V_inc, max_V)
                elif el['group'] == 'base':
                    max_M_base = max(max_M_base, max_M); max_V_base = max(max_V_base, max_V)
        
        struts_results = []
        for el in fea_data['elements']:
            if el['type'] == 'truss':
                struts_results.append({'type': el['sec'], 'N': abs(el['internal']['N'][0])})
            
        fea_data['sys_data'].update({
            'max_M_inc': max_M_inc, 'max_V_inc': max_V_inc,
            'max_M_base': max_M_base, 'max_V_base': max_V_base,
            'struts_res': struts_results,
            'img_bufs': img_bufs
        })
        
        docx_out = generate_inclined_report(fea_data['sys_data'])
        
        st.success("✅ SAP2000-Style Analysis Complete!")
        st.download_button("⬇️ Download Inclined System Calculation Sheet", 
                           data=docx_out.getvalue(), 
                           file_name="Inclined_System_Report.docx", 
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
